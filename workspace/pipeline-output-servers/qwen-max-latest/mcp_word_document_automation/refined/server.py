import json
import os
from docx import Document
from docx.shared import Inches, Pt
from mcp.server.fastmcp import FastMCP

# Initialize FastMCP server
mcp = FastMCP("mcp_word_document_automation")

class WordDocumentHandler:
    def __init__(self):
        self.document = None

    def create_document(self):
        """Create a new Word document.

        Returns:
            A string indicating the successful creation of the document.
        """
        try:
            self.document = Document()
            return json.dumps({"message": "Document created successfully."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def open_document(self, file_path: str):
        """Open an existing Word document.

        Args:
            file_path: (str) The path to the file to open.

        Returns:
            A string indicating the successful opening of the document.
        """
        try:
            if not os.path.isabs(file_path):
                file_path = os.path.abspath(file_path)
            if not os.path.exists(file_path):
                return json.dumps({"error": f"File not found: {file_path}"})
            self.document = Document(file_path)
            return json.dumps({"message": f"Document opened successfully from {file_path}."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def save_document(self):
        """Save the current Word document.

        Returns:
            A string indicating the successful saving of the document.
        """
        try:
            if self.document is not None:
                self.document.save('current_document.docx')
                return json.dumps({"message": "Document saved successfully."})
            else:
                return json.dumps({"error": "No document to save."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def save_as_document(self, file_path: str):
        """Save the current Word document with a new name or location.

        Args:
            file_path: (str) The new file path to save the document.

        Returns:
            A string indicating the successful saving of the document.
        """
        try:
            if self.document is not None:
                if not os.path.isabs(file_path):
                    file_path = os.path.abspath(file_path)
                self.document.save(file_path)
                return json.dumps({"message": f"Document saved successfully to {file_path}."})
            else:
                return json.dumps({"error": "No document to save."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def create_document_copy(self, file_path: str):
        """Create a copy of the current Word document.

        Args:
            file_path: (str) The destination path for the document copy.

        Returns:
            A string indicating the successful creation of the document copy.
        """
        try:
            if self.document is not None:
                if not os.path.isabs(file_path):
                    file_path = os.path.abspath(file_path)
                self.document.save(file_path)
                return json.dumps({"message": f"Document copy created successfully at {file_path}."})
            else:
                return json.dumps({"error": "No document to copy."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def add_paragraph(self, text: str):
        """Add a paragraph of text to the document.

        Args:
            text: (str) The text content of the paragraph.

        Returns:
            A string indicating the successful addition of the paragraph.
        """
        try:
            if self.document is not None:
                self.document.add_paragraph(text)
                return json.dumps({"message": "Paragraph added successfully."})
            else:
                return json.dumps({"error": "No document to add paragraph to."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def add_heading(self, text: str, level: int):
        """Add a heading to the document.

        Args:
            text: (str) The text of the heading.
            level: (int) The level of the heading (e.g., 1 for main heading).

        Returns:
            A string indicating the successful addition of the heading.
        """
        try:
            if self.document is not None:
                if not 1 <= level <= 9:
                    return json.dumps({"error": "Heading level must be between 1 and 9."})
                self.document.add_heading(text, level=level)
                return json.dumps({"message": "Heading added successfully."})
            else:
                return json.dumps({"error": "No document to add heading to."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def add_table(self, rows: int, cols: int):
        """Add a table to the document.

        Args:
            rows: (int) Number of rows in the table.
            cols: (int) Number of columns in the table.

        Returns:
            A string indicating the successful addition of the table.
        """
        try:
            if self.document is not None:
                if rows <= 0 or cols <= 0:
                    return json.dumps({"error": "Rows and columns must be greater than zero."})
                self.document.add_table(rows=rows, cols=cols)
                return json.dumps({"message": "Table added successfully."})
            else:
                return json.dumps({"error": "No document to add table to."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def add_page_break(self):
        """Add a page break to the document.

        Returns:
            A string indicating the successful addition of the page break.
        """
        try:
            if self.document is not None:
                self.document.add_page_break()
                return json.dumps({"message": "Page break added successfully."})
            else:
                return json.dumps({"error": "No document to add page break to."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def get_document_info(self):
        """Retrieve information about the current document.

        Returns:
            A dictionary containing various pieces of document information.
        """
        try:
            if self.document is not None:
                core_properties = self.document.core_properties
                info = {
                    "author": core_properties.author,
                    "category": core_properties.category,
                    "comments": core_properties.comments,
                    "content_status": core_properties.content_status,
                    "created": core_properties.created.isoformat() if core_properties.created else None,
                    "identifier": core_properties.identifier,
                    "keywords": core_properties.keywords,
                    "language": core_properties.language,
                    "last_modified_by": core_properties.last_modified_by,
                    "last_printed": core_properties.last_printed.isoformat() if core_properties.last_printed else None,
                    "modified": core_properties.modified.isoformat() if core_properties.modified else None,
                    "revision": core_properties.revision,
                    "subject": core_properties.subject,
                    "title": core_properties.title,
                    "version": core_properties.version
                }
                return json.dumps(info)
            else:
                return json.dumps({"error": "No document to get information from."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def search_text(self, search_string: str):
        """Search for specific text within the document.

        Args:
            search_string: (str) The text to search for.

        Returns:
            A list of positions where the text is found.
        """
        try:
            if self.document is not None:
                positions = []
                for i, paragraph in enumerate(self.document.paragraphs):
                    if search_string in paragraph.text:
                        positions.append(i)
                return json.dumps({"positions": positions})
            else:
                return json.dumps({"error": "No document to search in."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def find_and_replace(self, find_string: str, replace_string: str):
        """Find and replace specific text within the document.

        Args:
            find_string: (str) The text to find.
            replace_string: (str) The text to replace it with.

        Returns:
            A string indicating the number of replacements made.
        """
        try:
            if self.document is not None:
                count = 0
                for paragraph in self.document.paragraphs:
                    if find_string in paragraph.text:
                        paragraph.text = paragraph.text.replace(find_string, replace_string)
                        count += 1
                return json.dumps({"replacements": count})
            else:
                return json.dumps({"error": "No document to perform find and replace in."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def delete_paragraph(self, paragraph_index: int):
        """Delete a specified paragraph from the document.

        Args:
            paragraph_index: (int) The index of the paragraph to delete.

        Returns:
            A string indicating the successful deletion of the paragraph.
        """
        try:
            if self.document is not None:
                if 0 <= paragraph_index < len(self.document.paragraphs):
                    p = self.document.paragraphs[paragraph_index]
                    p._element.getparent().remove(p._element)
                    p._p = p._element = None
                    return json.dumps({"message": "Paragraph deleted successfully."})
                else:
                    return json.dumps({"error": "Invalid paragraph index."})
            else:
                return json.dumps({"error": "No document to delete paragraph from."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def delete_text(self, text: str):
        """Delete all instances of a specified text from the document.

        Args:
            text: (str) The text to delete.

        Returns:
            A string indicating the number of deletions made.
        """
        try:
            if self.document is not None:
                count = 0
                for paragraph in self.document.paragraphs:
                    if text in paragraph.text:
                        paragraph.text = paragraph.text.replace(text, '')
                        count += 1
                return json.dumps({"deletions": count})
            else:
                return json.dumps({"error": "No document to delete text from."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def add_table_row(self, table_index: int):
        """Add a row to a specified table in the document.

        Args:
            table_index: (int) The index of the table to add a row to.

        Returns:
            A string indicating the successful addition of the row.
        """
        try:
            if self.document is not None and 0 <= table_index < len(self.document.tables):
                table = self.document.tables[table_index]
                table.add_row()
                return json.dumps({"message": "Row added successfully."})
            else:
                return json.dumps({"error": "Invalid table index."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def delete_table_row(self, table_index: int, row_index: int):
        """Delete a specified row from a table in the document.

        Args:
            table_index: (int) The index of the table.
            row_index: (int) The index of the row to delete.

        Returns:
            A string indicating the successful deletion of the row.
        """
        try:
            if self.document is not None and 0 <= table_index < len(self.document.tables):
                table = self.document.tables[table_index]
                if 0 <= row_index < len(table.rows):
                    row = table.rows[row_index]
                    row._element.getparent().remove(row._element)
                    return json.dumps({"message": "Row deleted successfully."})
                else:
                    return json.dumps({"error": "Invalid row index."})
            else:
                return json.dumps({"error": "Invalid table index."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def edit_table_cell(self, table_index: int, row_index: int, col_index: int, new_content: str):
        """Edit the content of a specified cell in a table.

        Args:
            table_index: (int) The index of the table.
            row_index: (int) The index of the row.
            col_index: (int) The index of the column.
            new_content: (str) The new content for the cell.

        Returns:
            A string indicating the successful editing of the cell.
        """
        try:
            if self.document is not None and 0 <= table_index < len(self.document.tables):
                table = self.document.tables[table_index]
                if 0 <= row_index < len(table.rows) and 0 <= col_index < len(table.columns):
                    cell = table.cell(row_index, col_index)
                    cell.text = new_content
                    return json.dumps({"message": "Cell edited successfully."})
                else:
                    return json.dumps({"error": "Invalid row or column index."})
            else:
                return json.dumps({"error": "Invalid table index."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def merge_table_cells(self, table_index: int, start_row: int, end_row: int, start_col: int, end_col: int):
        """Merge specified cells in a table.

        Args:
            table_index: (int) The index of the table.
            start_row: (int) The starting row index for the merge.
            end_row: (int) The ending row index for the merge.
            start_col: (int) The starting column index for the merge.
            end_col: (int) The ending column index for the merge.

        Returns:
            A string indicating the successful merging of cells.
        """
        try:
            if self.document is not None and 0 <= table_index < len(self.document.tables):
                table = self.document.tables[table_index]
                start_cell = table.cell(start_row, start_col)
                end_cell = table.cell(end_row, end_col)
                start_cell.merge(end_cell)
                return json.dumps({"message": "Cells merged successfully."})
            else:
                return json.dumps({"error": "Invalid table index."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def split_table(self, table_index: int, row_index: int):
        """Split a table into two tables at a specified row.

        Args:
            table_index: (int) The index of the table to split.
            row_index: (int) The row index at which to split the table.

        Returns:
            A string indicating the successful splitting of the table.
        """
        try:
            if self.document is not None and 0 <= table_index < len(self.document.tables):
                table = self.document.tables[table_index]
                if 0 <= row_index < len(table.rows):
                    # Create a new table below the original one
                    new_table = self.document.add_table(rows=len(table.rows) - row_index, cols=len(table.columns))
                    for i, row in enumerate(table.rows[row_index:], start=0):
                        for j, cell in enumerate(row.cells, start=0):
                            new_table.cell(i, j).text = cell.text
                    # Remove rows from the original table
                    for row in reversed(table.rows[row_index:]):
                        row._element.getparent().remove(row._element)
                    return json.dumps({"message": "Table split successfully."})
                else:
                    return json.dumps({"error": "Invalid row index."})
            else:
                return json.dumps({"error": "Invalid table index."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def set_page_margins(self, top: float, right: float, bottom: float, left: float):
        """Set the margins for the pages in the document.

        Args:
            top: (float) The top margin size in inches.
            right: (float) The right margin size in inches.
            bottom: (float) The bottom margin size in inches.
            left: (float) The left margin size in inches.

        Returns:
            A string indicating the successful setting of page margins.
        """
        try:
            if self.document is not None:
                sections = self.document.sections
                for section in sections:
                    section.top_margin = Inches(top)
                    section.right_margin = Inches(right)
                    section.bottom_margin = Inches(bottom)
                    section.left_margin = Inches(left)
                return json.dumps({"message": "Page margins set successfully."})
            else:
                return json.dumps({"error": "No document to set page margins for."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def replace_section(self, heading: str, new_content: str):
        """Replace the content of a section based on the heading.

        Args:
            heading: (str) The heading of the section to replace.
            new_content: (str) The new content for the section.

        Returns:
            A string indicating the successful replacement of the section.
        """
        try:
            if self.document is not None:
                replaced = False
                for paragraph in self.document.paragraphs:
                    if paragraph.text.startswith(heading):
                        paragraph.text = new_content
                        replaced = True
                        break
                if replaced:
                    return json.dumps({"message": "Section replaced successfully."})
                else:
                    return json.dumps({"error": "Heading not found."})
            else:
                return json.dumps({"error": "No document to replace section in."})
        except Exception as e:
            return json.dumps({"error": str(e)})

    def edit_section_by_keyword(self, keyword: str, new_content: str):
        """Edit the content of a section based on a keyword.

        Args:
            keyword: (str) The keyword to locate the section.
            new_content: (str) The new content for the section.

        Returns:
            A string indicating the successful editing of the section.
        """
        try:
            if self.document is not None:
                edited = False
                for paragraph in self.document.paragraphs:
                    if keyword in paragraph.text:
                        paragraph.text = new_content
                        edited = True
                if edited:
                    return json.dumps({"message": "Section edited successfully."})
                else:
                    return json.dumps({"error": "Keyword not found."})
            else:
                return json.dumps({"error": "No document to edit section in."})
        except Exception as e:
            return json.dumps({"error": str(e)})

# Instantiate the handler
document_handler = WordDocumentHandler()

@mcp.tool()
def create_document():
    """Create a new Word document.

    Returns:
        A string indicating the successful creation of the document.
    """
    return document_handler.create_document()

@mcp.tool()
def open_document(file_path: str):
    """Open an existing Word document.

    Args:
        file_path: (str) The path to the file to open.

    Returns:
        A string indicating the successful opening of the document.
    """
    return document_handler.open_document(file_path)

@mcp.tool()
def save_document():
    """Save the current Word document.

    Returns:
        A string indicating the successful saving of the document.
    """
    return document_handler.save_document()

@mcp.tool()
def save_as_document(file_path: str):
    """Save the current Word document with a new name or location.

    Args:
        file_path: (str) The new file path to save the document.

    Returns:
        A string indicating the successful saving of the document.
    """
    return document_handler.save_as_document(file_path)

@mcp.tool()
def create_document_copy(file_path: str):
    """Create a copy of the current Word document.

    Args:
        file_path: (str) The destination path for the document copy.

    Returns:
        A string indicating the successful creation of the document copy.
    """
    return document_handler.create_document_copy(file_path)

@mcp.tool()
def add_paragraph(text: str):
    """Add a paragraph of text to the document.

    Args:
        text: (str) The text content of the paragraph.

    Returns:
        A string indicating the successful addition of the paragraph.
    """
    return document_handler.add_paragraph(text)

@mcp.tool()
def add_heading(text: str, level: int):
    """Add a heading to the document.

    Args:
        text: (str) The text of the heading.
        level: (int) The level of the heading (e.g., 1 for main heading).

    Returns:
        A string indicating the successful addition of the heading.
    """
    return document_handler.add_heading(text, level)

@mcp.tool()
def add_table(rows: int, cols: int):
    """Add a table to the document.

    Args:
        rows: (int) Number of rows in the table.
        cols: (int) Number of columns in the table.

    Returns:
        A string indicating the successful addition of the table.
    """
    return document_handler.add_table(rows, cols)

@mcp.tool()
def add_page_break():
    """Add a page break to the document.

    Returns:
        A string indicating the successful addition of the page break.
    """
    return document_handler.add_page_break()

@mcp.tool()
def get_document_info():
    """Retrieve information about the current document.

    Returns:
        A dictionary containing various pieces of document information.
    """
    return document_handler.get_document_info()

@mcp.tool()
def search_text(search_string: str):
    """Search for specific text within the document.

    Args:
        search_string: (str) The text to search for.

    Returns:
        A list of positions where the text is found.
    """
    return document_handler.search_text(search_string)

@mcp.tool()
def find_and_replace(find_string: str, replace_string: str):
    """Find and replace specific text within the document.

    Args:
        find_string: (str) The text to find.
        replace_string: (str) The text to replace it with.

    Returns:
        A string indicating the number of replacements made.
    """
    return document_handler.find_and_replace(find_string, replace_string)

@mcp.tool()
def delete_paragraph(paragraph_index: int):
    """Delete a specified paragraph from the document.

    Args:
        paragraph_index: (int) The index of the paragraph to delete.

    Returns:
        A string indicating the successful deletion of the paragraph.
    """
    return document_handler.delete_paragraph(paragraph_index)

@mcp.tool()
def delete_text(text: str):
    """Delete all instances of a specified text from the document.

    Args:
        text: (str) The text to delete.

    Returns:
        A string indicating the number of deletions made.
    """
    return document_handler.delete_text(text)

@mcp.tool()
def add_table_row(table_index: int):
    """Add a row to a specified table in the document.

    Args:
        table_index: (int) The index of the table to add a row to.

    Returns:
        A string indicating the successful addition of the row.
    """
    return document_handler.add_table_row(table_index)

@mcp.tool()
def delete_table_row(table_index: int, row_index: int):
    """Delete a specified row from a table in the document.

    Args:
        table_index: (int) The index of the table.
        row_index: (int) The index of the row to delete.

    Returns:
        A string indicating the successful deletion of the row.
    """
    return document_handler.delete_table_row(table_index, row_index)

@mcp.tool()
def edit_table_cell(table_index: int, row_index: int, col_index: int, new_content: str):
    """Edit the content of a specified cell in a table.

    Args:
        table_index: (int) The index of the table.
        row_index: (int) The index of the row.
        col_index: (int) The index of the column.
        new_content: (str) The new content for the cell.

    Returns:
        A string indicating the successful editing of the cell.
    """
    return document_handler.edit_table_cell(table_index, row_index, col_index, new_content)

@mcp.tool()
def merge_table_cells(table_index: int, start_row: int, end_row: int, start_col: int, end_col: int):
    """Merge specified cells in a table.

    Args:
        table_index: (int) The index of the table.
        start_row: (int) The starting row index for the merge.
        end_row: (int) The ending row index for the merge.
        start_col: (int) The starting column index for the merge.
        end_col: (int) The ending column index for the merge.

    Returns:
        A string indicating the successful merging of cells.
    """
    return document_handler.merge_table_cells(table_index, start_row, end_row, start_col, end_col)

@mcp.tool()
def split_table(table_index: int, row_index: int):
    """Split a table into two tables at a specified row.

    Args:
        table_index: (int) The index of the table to split.
        row_index: (int) The row index at which to split the table.

    Returns:
        A string indicating the successful splitting of the table.
    """
    return document_handler.split_table(table_index, row_index)

@mcp.tool()
def set_page_margins(top: float, right: float, bottom: float, left: float):
    """Set the margins for the pages in the document.

    Args:
        top: (float) The top margin size in inches.
        right: (float) The right margin size in inches.
        bottom: (float) The bottom margin size in inches.
        left: (float) The left margin size in inches.

    Returns:
        A string indicating the successful setting of page margins.
    """
    return document_handler.set_page_margins(top, right, bottom, left)

@mcp.tool()
def replace_section(heading: str, new_content: str):
    """Replace the content of a section based on the heading.

    Args:
        heading: (str) The heading of the section to replace.
        new_content: (str) The new content for the section.

    Returns:
        A string indicating the successful replacement of the section.
    """
    return document_handler.replace_section(heading, new_content)

@mcp.tool()
def edit_section_by_keyword(keyword: str, new_content: str):
    """Edit the content of a section based on a keyword.

    Args:
        keyword: (str) The keyword to locate the section.
        new_content: (str) The new content for the section.

    Returns:
        A string indicating the successful editing of the section.
    """
    return document_handler.edit_section_by_keyword(keyword, new_content)

if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding='utf-8-sig')
    mcp.run()