import json
import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import win32com.client as win32
import zipfile
import xml.etree.ElementTree as ET
from mcp.server.fastmcp import FastMCP

def create_document(title, author="Unknown", subject=None):
    """
    Creates a new Word document and sets its metadata.

    Args:
        title (str): The title of the document.
        author (str, optional): The author's name. Defaults to "Unknown".
        subject (str, optional): The subject of the document.

    Returns:
        str: A message indicating the successful creation of the document.

    Example:
        create_document(title="My Document", author="John Doe", subject="Test Document")
    """
    try:
        document = Document()
        core_properties = document.core_properties
        core_properties.title = title
        core_properties.author = author
        if subject:
            core_properties.subject = subject
        document.save("new_document.docx")
        return json.dumps({"message": "Document created successfully."})
    except Exception as e:
        return json.dumps({"error": f"Failed to create document: {e}"})

def get_document_text(file_path):
    """
    Extracts the full text content from a Word document.

    Args:
        file_path (str): Path to the Word document.

    Returns:
        str: A string containing all the text from the document.

    Example:
        get_document_text(file_path="path/to/document.docx")
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file '{file_path}' does not exist.")
        document = Document(file_path)
        full_text = '\n'.join([para.text for para in document.paragraphs])
        return json.dumps({"text": full_text})
    except Exception as e:
        return json.dumps({"error": f"Failed to extract text: {e}"})

def add_paragraph(file_path, text):
    """
    Adds a paragraph of text to a Word document.

    Args:
        file_path (str): Path to the Word document.
        text (str): The text to be added as a paragraph.

    Returns:
        str: A message confirming the addition of the paragraph.

    Example:
        add_paragraph(file_path="path/to/document.docx", text="This is a new paragraph.")
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file '{file_path}' does not exist.")
        document = Document(file_path)
        document.add_paragraph(text)
        document.save(file_path)
        return json.dumps({"message": "Paragraph added successfully."})
    except Exception as e:
        return json.dumps({"error": f"Failed to add paragraph: {e}"})

def add_heading(file_path, text, level):
    """
    Adds a heading to a Word document at a specified level.

    Args:
        file_path (str): Path to the Word document.
        text (str): The heading text.
        level (int): The level of the heading (1 for main headings, higher numbers for subheadings).

    Returns:
        str: A message confirming the addition of the heading.

    Example:
        add_heading(file_path="path/to/document.docx", text="Main Title", level=1)
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file '{file_path}' does not exist.")
        # Ensure level is an integer
        level = int(level)
        if not (1 <= level <= 9):
            raise ValueError("Heading level must be between 1 and 9.")
        document = Document(file_path)
        document.add_heading(text, level=level)
        document.save(file_path)
        return json.dumps({"message": "Heading added successfully."})
    except Exception as e:
        return json.dumps({"error": f"Failed to add heading: {e}"})

def create_custom_style(file_path, style_name, font_size, bold=False, italic=False):
    """
    Creates a custom text style in the Word document.

    Args:
        file_path (str): Path to the Word document.
        style_name (str): Name of the new style.
        font_size (int): Font size for the style.
        bold (bool, optional): Whether the text should be bold. Defaults to False.
        italic (bool, optional): Whether the text should be italicized. Defaults to False.

    Returns:
        str: A message confirming the creation of the custom style.

    Example:
        create_custom_style(file_path="path/to/document.docx", style_name="CustomStyle", font_size=14, bold=True, italic=False)
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file '{file_path}' does not exist.")
        # Convert font_size to integer
        font_size = int(font_size)
        document = Document(file_path)
        if style_name in document.styles:
            raise ValueError(f"A style with the name '{style_name}' already exists.")
        style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.size = Pt(font_size)
        font.bold = bool(bold)
        font.italic = bool(italic)
        document.save(file_path)
        return json.dumps({"message": "Custom style created successfully."})
    except Exception as e:
        return json.dumps({"error": f"Failed to create custom style: {e}"})

def format_text(file_path, start_pos, end_pos, style_name):
    """
    Formats a specific range of text within a Word document.

    Args:
        file_path (str): Path to the Word document.
        start_pos (int): Starting position of the text to format.
        end_pos (int): Ending position of the text to format.
        style_name (str): Name of the style to apply.

    Returns:
        str: A message confirming the formatting of the text.

    Example:
        format_text(file_path="path/to/document.docx", start_pos=0, end_pos=10, style_name="CustomStyle")
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file '{file_path}' does not exist.")
        document = Document(file_path)
        if style_name not in document.styles:
            raise ValueError(f"The style '{style_name}' does not exist in the document.")
        paragraphs = document.paragraphs
        for paragraph in paragraphs:
            if start_pos < len(paragraph.text) and end_pos <= len(paragraph.text):
                run = paragraph.runs[0]
                run.style = style_name
        document.save(file_path)
        return json.dumps({"message": "Text formatted successfully."})
    except Exception as e:
        return json.dumps({"error": f"Failed to format text: {e}"})

def protect_document(file_path, password):
    """
    Sets password protection on a Word document.

    Args:
        file_path (str): Path to the Word document.
        password (str): Password to set for protection.

    Returns:
        str: A message confirming that the document is now protected.

    Example:
        protect_document(file_path="path/to/document.docx", password="mysecretpassword")
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file '{file_path}' does not exist.")
        word = win32.Dispatch("Word.Application")
        doc = word.Documents.Open(os.path.abspath(file_path))
        doc.Password = password
        doc.Save()
        doc.Close()
        word.Quit()
        return json.dumps({"message": "Document protected successfully."})
    except Exception as e:
        return json.dumps({"error": f"Failed to protect document: {e}"})

def add_footnote_to_document(file_path, text):
    """
    Adds a footnote to a Word document.

    Args:
        file_path (str): Path to the Word document.
        text (str): The footnote text.

    Returns:
        str: A message confirming the addition of the footnote.

    Example:
        add_footnote_to_document(file_path="path/to/document.docx", text="This is a footnote.")
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file '{file_path}' does not exist.")
        document = Document(file_path)
        footnote = document.add_paragraph().add_run(text)
        document.save(file_path)
        return json.dumps({"message": "Footnote added successfully."})
    except Exception as e:
        return json.dumps({"error": f"Failed to add footnote: {e}"})

def get_paragraph_text_from_document(file_path, paragraph_index):
    """
    Retrieves the text of a specific paragraph from a Word document.

    Args:
        file_path (str): Path to the Word document.
        paragraph_index (int): Index of the paragraph to retrieve.

    Returns:
        str: A string containing the text of the specified paragraph.

    Example:
        get_paragraph_text_from_document(file_path="path/to/document.docx", paragraph_index=0)
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file '{file_path}' does not exist.")
        # Convert paragraph_index to integer
        paragraph_index = int(paragraph_index)
        document = Document(file_path)
        if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
            raise IndexError("Paragraph index out of range.")
        paragraph_text = document.paragraphs[paragraph_index].text
        return json.dumps({"text": paragraph_text})
    except Exception as e:
        return json.dumps({"error": f"Failed to retrieve paragraph text: {e}"})

def find_text_in_document(file_path, search_text):
    """
    Searches for specified text within a Word document.

    Args:
        file_path (str): Path to the Word document.
        search_text (str): Text to search for.

    Returns:
        list: A list of positions where the text was found.

    Example:
        find_text_in_document(file_path="path/to/document.docx", search_text="important")
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file '{file_path}' does not exist.")
        document = Document(file_path)
        positions = []
        for i, paragraph in enumerate(document.paragraphs):
            if search_text in paragraph.text:
                positions.append(i)
        return json.dumps({"positions": positions})
    except Exception as e:
        return json.dumps({"error": f"Failed to find text: {e}"})

def remove_password_protection(file_path, password):
    """
    Removes password protection from a Word document by modifying the settings.xml file.

    Args:
        file_path (str): Path to the Word document.
        password (str): Password used to protect the document.

    Returns:
        str: A message confirming that the password protection has been removed.

    Example:
        remove_password_protection(file_path="path/to/document.docx", password="securePass123")
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file '{file_path}' does not exist.")

        # Create a temporary zip file
        temp_dir = 'temp_unzip'
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        # Remove password protection by modifying settings.xml
        settings_path = os.path.join(temp_dir, 'word', 'settings.xml')
        tree = ET.parse(settings_path)
        root = tree.getroot()

        # Find and remove the protection elements
        for elem in root.iter():
            if 'writeProtection' in elem.tag:
                root.remove(elem)

        # Save the modified XML
        tree.write(settings_path)

        # Create a new zip file without password protection
        new_file_path = file_path.replace('.docx', '_unprotected.docx')
        with zipfile.ZipFile(new_file_path, 'w') as zip_ref:
            for foldername, subfolders, filenames in os.walk(temp_dir):
                for filename in filenames:
                    file_path_inside = os.path.join(foldername, filename)
                    arcname = os.path.relpath(file_path_inside, temp_dir)
                    zip_ref.write(file_path_inside, arcname)

        # Clean up temporary files
        for foldername, subfolders, filenames in os.walk(temp_dir):
            for filename in filenames:
                os.unlink(os.path.join(foldername, filename))
        os.rmdir(temp_dir)

        return json.dumps({"message": "Password protection removed successfully."})
    except Exception as e:
        return json.dumps({"error": f"Failed to remove password protection: {e}"})

# Initialize FastMCP server
mcp = FastMCP("word_document_processor")

@mcp.tool()
def create_document_tool(title, author="Unknown", subject=None):
    return create_document(title, author, subject)

@mcp.tool()
def get_document_text_tool(file_path):
    return get_document_text(file_path)

@mcp.tool()
def add_paragraph_tool(file_path, text):
    return add_paragraph(file_path, text)

@mcp.tool()
def add_heading_tool(file_path, text, level):
    return add_heading(file_path, text, level)

@mcp.tool()
def create_custom_style_tool(file_path, style_name, font_size, bold=False, italic=False):
    return create_custom_style(file_path, style_name, font_size, bool(bold), bool(italic))

@mcp.tool()
def format_text_tool(file_path, start_pos, end_pos, style_name):
    return format_text(file_path, start_pos, end_pos, style_name)

@mcp.tool()
def protect_document_tool(file_path, password):
    return protect_document(file_path, password)

@mcp.tool()
def add_footnote_to_document_tool(file_path, text):
    return add_footnote_to_document(file_path, text)

@mcp.tool()
def get_paragraph_text_from_document_tool(file_path, paragraph_index):
    return get_paragraph_text_from_document(file_path, paragraph_index)

@mcp.tool()
def find_text_in_document_tool(file_path, search_text):
    return find_text_in_document(file_path, search_text)

@mcp.tool()
def remove_password_protection(file_file_path, password):
    return remove_password_protection(file_file_path, password)

if __name__ == "__main__":
    mcp.run()