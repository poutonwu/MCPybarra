import sys
import os
import json
from mcp.server.fastmcp import FastMCP
import docx
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from pdf2docx import Converter

# Initialize FastMCP server
mcp = FastMCP("mcp_word_document_automation")

def _validate_file_path(file_path: str, must_exist: bool = True):
    """Helper function to validate file paths."""
    if not file_path or not isinstance(file_path, str):
        raise ValueError("File path must be a non-empty string.")
    
    abs_path = os.path.abspath(file_path)

    if must_exist and not os.path.exists(abs_path):
        raise FileNotFoundError(f"The specified file does not exist: {abs_path}")
    
    if not abs_path.lower().endswith('.docx'):
        raise ValueError("File path must point to a .docx file.")
        
    return abs_path

@mcp.tool()
def create_document(file_path: str, title: str = None, author: str = None, subject: str = None, keywords: str = None) -> dict:
    """
    Creates a new, empty Word document (.docx) and sets its core metadata properties.

    Args:
        file_path (str): The full path where the new Word document will be saved (e.g., "C:/docs/mydocument.docx").
        title (str, optional): The title of the document.
        author (str, optional): The author's name.
        subject (str, optional): The subject of the document.
        keywords (str, optional): Comma-separated keywords for the document.

    Returns:
        dict: A dictionary confirming the action, e.g., {"status": "success", "file_path": "C:/docs/mydocument.docx"}.

    Example:
        create_document(file_path="mydocument.docx", title="My Doc", author="John Doe")
    """
    try:
        # Validate that the path is for a .docx file, but it doesn't have to exist yet.
        if not file_path or not isinstance(file_path, str) or not file_path.lower().endswith('.docx'):
            raise ValueError("File path must be a non-empty string ending with .docx")
        
        document = docx.Document()
        core_properties = document.core_properties
        if title: core_properties.title = title
        if author: core_properties.author = author
        if subject: core_properties.subject = subject
        if keywords: core_properties.keywords = keywords
        
        document.save(file_path)
        return {"status": "success", "file_path": os.path.abspath(file_path)}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@mcp.tool()
def get_document_text(file_path: str) -> str:
    """
    Extracts and returns all text content from an existing Word document.

    Args:
        file_path (str): The path to the Word document to be read.

    Returns:
        str: The complete text content of the document, with paragraphs separated by newline characters.

    Example:
        get_document_text(file_path="mydocument.docx")
    """
    try:
        path = _validate_file_path(file_path)
        document = docx.Document(path)
        full_text = [para.text for para in document.paragraphs]
        return '\n'.join(full_text)
    except (ValueError, FileNotFoundError) as e:
        return json.dumps({"status": "error", "message": str(e)})
    except Exception as e:
        return json.dumps({"status": "error", "message": f"An unexpected error occurred: {e}"})

@mcp.tool()
def add_paragraph(file_path: str, text: str) -> dict:
    """
    Adds a new paragraph of text to the end of a specified Word document.

    Args:
        file_path (str): The path to the Word document.
        text (str): The text content to be added as a new paragraph.

    Returns:
        dict: A dictionary confirming the action.

    Example:
        add_paragraph(file_path="mydocument.docx", text="This is a new paragraph.")
    """
    try:
        path = _validate_file_path(file_path)
        document = docx.Document(path)
        document.add_paragraph(text)
        document.save(path)
        return {"status": "success", "message": "Paragraph added."}
    except (ValueError, FileNotFoundError) as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"An unexpected error occurred: {e}"}

@mcp.tool()
def add_heading(file_path: str, text: str, level: int) -> dict:
    """
    Adds a heading to the end of a Word document with a specified level.

    Args:
        file_path (str): The path to the Word document.
        text (str): The text of the heading.
        level (int): The heading level, from 0 (Title) to 9. Level 1 is a main heading.

    Returns:
        dict: A dictionary confirming the action.

    Example:
        add_heading(file_path="mydocument.docx", text="Chapter 1", level=1)
    """
    try:
        if not isinstance(level, int) or not (0 <= level <= 9):
            raise ValueError("Heading level must be an integer between 0 and 9.")
        path = _validate_file_path(file_path)
        document = docx.Document(path)
        document.add_heading(text, level=level)
        document.save(path)
        return {"status": "success", "message": "Heading added."}
    except (ValueError, FileNotFoundError) as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"An unexpected error occurred: {e}"}

@mcp.tool()
def create_custom_style(file_path: str, style_name: str, font_name: str = "Calibri", font_size_pt: int = 12, bold: bool = False, italic: bool = False) -> dict:
    """
    Creates a new custom paragraph style within the document.

    Args:
        file_path (str): The path to the Word document.
        style_name (str): The name for the new custom style.
        font_name (str, optional): The font name (e.g., "Calibri", "Times New Roman").
        font_size_pt (int, optional): The font size in points (e.g., 12).
        bold (bool, optional): Whether the text should be bold. Defaults to False.
        italic (bool, optional): Whether the text should be italic. Defaults to False.

    Returns:
        dict: A dictionary confirming the action.

    Example:
        create_custom_style(file_path="mydocument.docx", style_name="MyStyle", font_name="Arial", font_size_pt=14, bold=True)
    """
    try:
        path = _validate_file_path(file_path)
        document = docx.Document(path)
        styles = document.styles
        
        if style_name in styles:
            return {"status": "error", "message": f"Style '{style_name}' already exists."}

        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        if font_name:
            font.name = font_name
        if font_size_pt:
            font.size = Pt(font_size_pt)
        font.bold = bold
        font.italic = italic
        
        document.save(path)
        return {"status": "success", "message": f"Style '{style_name}' created."}
    except (ValueError, FileNotFoundError) as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"An unexpected error occurred: {e}"}

@mcp.tool()
def format_text(file_path: str, search_text: str, bold: bool = False, italic: bool = False) -> dict:
    """
    Finds the first occurrence of a specific text string in the document and applies bold and/or italic formatting to it.
    This is a simplified implementation. It formats the entire run containing the text, not just the text itself.

    Args:
        file_path (str): The path to the Word document.
        search_text (str): The text to find and format.
        bold (bool, optional): Apply bold formatting. Defaults to False.
        italic (bool, optional): Apply italic formatting. Defaults to False.

    Returns:
        dict: A dictionary indicating success or failure.

    Example:
        format_text(file_path="mydocument.docx", search_text="important", bold=True)
    """
    try:
        path = _validate_file_path(file_path)
        document = docx.Document(path)
        found = False
        for paragraph in document.paragraphs:
            if search_text in paragraph.text:
                for run in paragraph.runs:
                    if search_text in run.text:
                        run.bold = bold
                        run.italic = italic
                        found = True
                        break # Format first occurrence in paragraph
            if found:
                break
        
        if found:
            document.save(path)
            return {"status": "success", "message": "Text formatted."}
        else:
            return {"status": "not_found", "message": "Text not found."}
            
    except (ValueError, FileNotFoundError) as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"An unexpected error occurred: {e}"}

@mcp.tool()
def protect_document(file_path: str, password: str) -> dict:
    """
    Sets write protection on a Word document. Note: This is not supported by the current library.

    Args:
        file_path (str): The path to the Word document.
        password (str): The password required to disable protection.

    Returns:
        dict: A dictionary confirming the attempt.

    Example:
        protect_document(file_path="mydocument.docx", password="password123")
    """
    return {"status": "error", "message": "Feature not supported by the current environment. 'python-docx' does not support password protection."}

@mcp.tool()
def add_footnote_to_document(file_path: str, paragraph_index: int, footnote_text: str) -> dict:
    """
    Adds a footnote to a specific paragraph. Note: This is not directly supported by the current library.

    Args:
        file_path (str): The path to the Word document.
        paragraph_index (int): The 0-based index of the paragraph for the footnote.
        footnote_text (str): The text content of the footnote.

    Returns:
        dict: A dictionary confirming the action.

    Example:
        add_footnote_to_document(file_path="mydocument.docx", paragraph_index=0, footnote_text="This is a footnote.")
    """
    return {"status": "error", "message": "Failed to add footnote due to library limitations. 'python-docx' does not have a direct API for footnotes."}

@mcp.tool()
def get_paragraph_text_from_document(file_path: str, paragraph_index: int) -> str:
    """
    Retrieves the text from a single, specific paragraph based on its index.

    Args:
        file_path (str): The path to the Word document.
        paragraph_index (int): The 0-based index of the paragraph to retrieve.

    Returns:
        str: The text content of the specified paragraph or an error message.

    Example:
        get_paragraph_text_from_document(file_path="mydocument.docx", paragraph_index=0)
    """
    try:
        path = _validate_file_path(file_path)
        document = docx.Document(path)
        if not isinstance(paragraph_index, int) or paragraph_index < 0:
             raise ValueError("Paragraph index must be a non-negative integer.")
        if paragraph_index >= len(document.paragraphs):
            return json.dumps({"status": "error", "message": "Paragraph index out of bounds."})
        return document.paragraphs[paragraph_index].text
    except (ValueError, FileNotFoundError) as e:
        return json.dumps({"status": "error", "message": str(e)})
    except Exception as e:
        return json.dumps({"status": "error", "message": f"An unexpected error occurred: {e}"})

@mcp.tool()
def find_text_in_document(file_path: str, search_text: str) -> list:
    """
    Searches the document for a text string and returns the indices of all paragraphs containing the text.

    Args:
        file_path (str): The path to the Word document.
        search_text (str): The text to search for.

    Returns:
        list: A list of integer indices for each paragraph where the text was found, or a list containing an error dictionary.

    Example:
        find_text_in_document(file_path="mydocument.docx", search_text="hello")
    """
    try:
        path = _validate_file_path(file_path)
        document = docx.Document(path)
        found_indices = [i for i, p in enumerate(document.paragraphs) if search_text in p.text]
        return found_indices
    except (ValueError, FileNotFoundError) as e:
        return [{"status": "error", "message": str(e)}]
    except Exception as e:
        return [{"status": "error", "message": f"An unexpected error occurred: {e}"}]

@mcp.tool()
def add_table(file_path: str, rows: int, cols: int) -> dict:
    """
    Adds a table with a specified number of rows and columns to the end of the document.

    Args:
        file_path (str): The path to the Word document.
        rows (int): The number of rows in the table.
        cols (int): The number of columns in the table.

    Returns:
        dict: A dictionary confirming the action.

    Example:
        add_table(file_path="mydocument.docx", rows=3, cols=4)
    """
    try:
        if not isinstance(rows, int) or rows <= 0 or not isinstance(cols, int) or cols <= 0:
            raise ValueError("Rows and columns must be positive integers.")
        path = _validate_file_path(file_path)
        document = docx.Document(path)
        document.add_table(rows=rows, cols=cols, style='Table Grid')
        document.save(path)
        return {"status": "success", "message": "Table added."}
    except (ValueError, FileNotFoundError) as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"An unexpected error occurred: {e}"}

@mcp.tool()
def add_image(doc_file_path: str, image_file_path: str, width_inches: float = None) -> dict:
    """
    Adds an image from a local path to the end of the document.

    Args:
        doc_file_path (str): The path to the Word document.
        image_file_path (str): The path to the image file to be inserted.
        width_inches (float, optional): The width of the image in inches.

    Returns:
        dict: A dictionary confirming the action.

    Example:
        add_image(doc_file_path="mydocument.docx", image_file_path="myimage.png", width_inches=2.5)
    """
    try:
        doc_path = _validate_file_path(doc_file_path)
        if not os.path.exists(image_file_path):
            raise FileNotFoundError(f"The specified image file does not exist: {image_file_path}")
            
        document = docx.Document(doc_path)
        if width_inches:
            document.add_picture(image_file_path, width=Inches(width_inches))
        else:
            document.add_picture(image_file_path)
        document.save(doc_path)
        return {"status": "success", "message": "Image added."}
    except (ValueError, FileNotFoundError) as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"An unexpected error occurred: {e}"}

@mcp.tool()
def add_page_break(file_path: str) -> dict:
    """
    Inserts a page break at the end of the document.

    Args:
        file_path (str): The path to the Word document.

    Returns:
        dict: A dictionary confirming the action.

    Example:
        add_page_break(file_path="mydocument.docx")
    """
    try:
        path = _validate_file_path(file_path)
        document = docx.Document(path)
        document.add_page_break()
        document.save(path)
        return {"status": "success", "message": "Page break added."}
    except (ValueError, FileNotFoundError) as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"An unexpected error occurred: {e}"}

@mcp.tool()
def add_footer(file_path: str, footer_text: str) -> dict:
    """
    Adds or replaces the footer for the default section of the document.

    Args:
        file_path (str): The path to the Word document.
        footer_text (str): The text to place in the footer.

    Returns:
        dict: A dictionary confirming the action.

    Example:
        add_footer(file_path="mydocument.docx", footer_text="Copyright 2024")
    """
    try:
        path = _validate_file_path(file_path)
        document = docx.Document(path)
        section = document.sections[0]
        footer = section.footer
        # Clear existing paragraphs in the footer before adding new text
        for p in footer.paragraphs:
            p.clear()
        footer.paragraphs[0].text = footer_text
        document.save(path)
        return {"status": "success", "message": "Footer updated."}
    except (ValueError, FileNotFoundError) as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"An unexpected error occurred: {e}"}

@mcp.tool()
def add_header(file_path: str, header_text: str) -> dict:
    """
    Adds or replaces the header for the default section of the document.

    Args:
        file_path (str): The path to the Word document.
        header_text (str): The text to place in the header.

    Returns:
        dict: A dictionary confirming the action.

    Example:
        add_header(file_path="mydocument.docx", header_text="My Document Header")
    """
    try:
        path = _validate_file_path(file_path)
        document = docx.Document(path)
        section = document.sections[0]
        header = section.header
        # Clear existing paragraphs in the header before adding new text
        for p in header.paragraphs:
            p.clear()
        header.paragraphs[0].text = header_text
        document.save(path)
        return {"status": "success", "message": "Header updated."}
    except (ValueError, FileNotFoundError) as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"An unexpected error occurred: {e}"}

@mcp.tool()
def convert_to_pdf(docx_path: str, pdf_path: str) -> dict:
    """
    Converts a Word document to a PDF file.

    Args:
        docx_path (str): The path to the source Word document.
        pdf_path (str): The path where the output PDF file will be saved.

    Returns:
        dict: A dictionary confirming the action.

    Example:
        convert_to_pdf(docx_path="mydocument.docx", pdf_path="mydocument.pdf")
    """
    try:
        doc_path = _validate_file_path(docx_path)
        if not pdf_path.lower().endswith('.pdf'):
            raise ValueError("Output path must be a .pdf file.")
            
        cv = Converter(doc_path)
        cv.convert(pdf_path, start=0, end=None)
        cv.close()
        
        if os.path.exists(pdf_path):
            return {"status": "success", "file_path": os.path.abspath(pdf_path)}
        else:
            return {"status": "error", "message": "PDF conversion failed for an unknown reason."}
    except (ValueError, FileNotFoundError) as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"An unexpected error occurred during PDF conversion: {e}"}

if __name__ == "__main__":
    if sys.platform == "win32":
        sys.stdout.reconfigure(encoding='utf-8')
    mcp.run()