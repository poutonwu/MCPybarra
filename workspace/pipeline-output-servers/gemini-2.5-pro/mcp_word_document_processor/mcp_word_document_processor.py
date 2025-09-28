import sys
import os
import json
import re
from mcp.server.fastmcp import FastMCP
import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx2pdf

# It's recommended to run Word and PDF conversion-related tasks in a separate process
# or thread in a production environment to avoid blocking the main application.
# For simplicity, this example runs them in the same process.

# Initialize FastMCP server
mcp = FastMCP("mcp_word_document_processor")

def _get_abs_path(file_path):
    """Converts a relative path to an absolute path and validates it."""
    if not isinstance(file_path, str) or not file_path.strip():
        raise ValueError("File path cannot be empty.")
    # Basic security check to prevent path traversal
    if ".." in file_path:
        raise ValueError("Invalid file path. Path traversal is not allowed.")
    return os.path.abspath(file_path)

@mcp.tool()
def create_document(filename: str, author: str = None, title: str = None, subject: str = None, keywords: str = None) -> str:
    """
    Creates a new, empty Word document with specified metadata and saves it.

    Args:
        filename (str): The name of the file to create (e.g., 'report.docx').
        author (str, optional): The author of the document. Defaults to None.
        title (str, optional): The title of the document. Defaults to None.
        subject (str, optional): The subject of the document. Defaults to None.
        keywords (str, optional): Comma-separated keywords for the document. Defaults to None.

    Returns:
        str: A JSON string confirming the successful creation, e.g.,
             '{"status": "success", "file_path": "report.docx"}'.

    Example:
        create_document(filename='new_report.docx', author='John Doe', title='Annual Report')
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document()

        core_properties = document.core_properties
        if author:
            core_properties.author = author
        if title:
            core_properties.title = title
        if subject:
            core_properties.subject = subject
        if keywords:
            core_properties.keywords = keywords

        document.save(abs_path)
        return json.dumps({'status': 'success', 'file_path': abs_path})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def get_document_text(filename: str) -> str:
    """
    Extracts and returns all the text content from an existing Word document.

    Args:
        filename (str): The path to the Word document.

    Returns:
        str: The full text content of the document.

    Example:
        get_document_text(filename='existing_report.docx')
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        full_text = [para.text for para in document.paragraphs]
        return '\n'.join(full_text)
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def add_paragraph(filename: str, text: str, style: str = None) -> str:
    """
    Adds a new paragraph of text to the end of a specified document.

    Args:
        filename (str): The path to the Word document.
        text (str): The text content to add in the paragraph.
        style (str, optional): The name of the style to apply (e.g., 'Normal'). Defaults to None.

    Returns:
        str: A JSON string confirming the operation.

    Example:
        add_paragraph(filename='report.docx', text='This is a new paragraph.')
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        document.add_paragraph(text, style=style)
        document.save(abs_path)
        return json.dumps({'status': 'success'})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def add_heading(filename: str, text: str, level: int) -> str:
    """
    Adds a heading to the document.

    Args:
        filename (str): The path to the Word document.
        text (str): The text of the heading.
        level (int): The heading level, from 0 (Title) to 9.

    Returns:
        str: A JSON string confirming the operation.

    Example:
        add_heading(filename='report.docx', text='Chapter 1', level=1)
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        document.add_heading(text, level=level)
        document.save(abs_path)
        return json.dumps({'status': 'success'})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def create_custom_style(filename: str, style_name: str, font_name: str = None, font_size_pt: int = None, bold: bool = False, italic: bool = False, underline: bool = False) -> str:
    """
    Creates a new custom paragraph style.

    Args:
        filename (str): Path to the Word document.
        style_name (str): Name for the new custom style.
        font_name (str, optional): Font for the style (e.g., 'Calibri').
        font_size_pt (int, optional): Font size in points.
        bold (bool, optional): Whether the font should be bold.
        italic (bool, optional): Whether the font should be italic.
        underline (bool, optional): Whether the font should be underlined.

    Returns:
        str: A JSON string confirming style creation.

    Example:
        create_custom_style(filename='report.docx', style_name='CustomBody', font_name='Arial', font_size_pt=12)
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        styles = document.styles
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        if font_name:
            font.name = font_name
        if font_size_pt:
            font.size = Pt(font_size_pt)
        font.bold = bold
        font.italic = italic
        font.underline = underline
        document.save(abs_path)
        return json.dumps({'status': 'success'})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def format_text(filename: str, search_text: str, bold: bool = False, italic: bool = False, underline: bool = False) -> str:
    """
    Searches for text and applies formatting to all instances found.

    Args:
        filename (str): Path to the Word document.
        search_text (str): The text to find and format.
        bold (bool, optional): Apply bold formatting.
        italic (bool, optional): Apply italic formatting.
        underline (bool, optional): Apply underline formatting.

    Returns:
        str: JSON string reporting the number of instances formatted.

    Example:
        format_text(filename='report.docx', search_text='important', bold=True)
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        count = 0
        for p in document.paragraphs:
            if search_text in p.text:
                # This is a simplified implementation. For complex cases,
                # a run-level search and replace would be needed to preserve other formatting.
                for run in p.runs:
                    if search_text in run.text:
                        # A more robust solution would split the run.
                        # This is a basic implementation for demonstration.
                        run.bold = bold or run.bold
                        run.italic = italic or run.italic
                        run.underline = underline or run.underline
                        count += 1
        document.save(abs_path)
        return json.dumps({'status': 'success', 'instances_formatted': count})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def protect_document(filename: str, password: str) -> str:
    """
    Sets write-protection for a document. Not supported by python-docx.

    Args:
        filename (str): Path to the Word document.
        password (str): The password for protection.

    Returns:
        str: A JSON string indicating the feature is not supported.

    Example:
        protect_document(filename='report.docx', password='123')
    """
    return json.dumps({'status': 'error', 'message': 'Password protection is not supported by the underlying library.'})

@mcp.tool()
def add_footnote_to_document(filename: str, footnote_text: str) -> str:
    """
    Adds a footnote to the end of the last paragraph.

    Args:
        filename (str): Path to the Word document.
        footnote_text (str): The text content of the footnote.

    Returns:
        str: A JSON string confirming the operation.

    Example:
        add_footnote_to_document(filename='report.docx', footnote_text='Source: Annual Report')
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        if not document.paragraphs:
            return json.dumps({'status': 'error', 'message': 'Document has no paragraphs to add a footnote to.'})

        last_paragraph = document.paragraphs[-1]
        last_paragraph.add_run().add_footnote(footnote_text)
        document.save(abs_path)
        return json.dumps({'status': 'success'})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def get_paragraph_text_from_document(filename: str, paragraph_index: int) -> str:
    """
    Retrieves the text from a single paragraph by its index.

    Args:
        filename (str): The path to the Word document.
        paragraph_index (int): The zero-based index of the paragraph.

    Returns:
        str: The text content of the specified paragraph.

    Example:
        get_paragraph_text_from_document(filename='report.docx', paragraph_index=0)
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        if 0 <= paragraph_index < len(document.paragraphs):
            return document.paragraphs[paragraph_index].text
        else:
            return json.dumps({'status': 'error', 'message': 'Paragraph index out of range.'})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def find_text_in_document(filename: str, search_text: str) -> str:
    """
    Searches for text and returns the content of all paragraphs containing it.

    Args:
        filename (str): The path to the Word document.
        search_text (str): The text to search for.

    Returns:
        str: A JSON string containing a list of paragraph texts.

    Example:
        find_text_in_document(filename='report.docx', search_text='budget')
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        found_paragraphs = [p.text for p in document.paragraphs if search_text in p.text]
        return json.dumps(found_paragraphs)
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def add_table(filename: str, rows: int, cols: int, style: str = 'Table Grid') -> str:
    """
    Adds a table with a specified number of rows and columns.

    Args:
        filename (str): Path to the Word document.
        rows (int): The number of rows in the table.
        cols (int): The number of columns in the table.
        style (str, optional): Style for the table (e.g., 'Light Shading Accent 1').

    Returns:
        str: A JSON string confirming the operation.

    Example:
        add_table(filename='report.docx', rows=3, cols=4)
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        document.add_table(rows=rows, cols=cols, style=style)
        document.save(abs_path)
        return json.dumps({'status': 'success'})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def add_picture(filename: str, picture_path: str, width_cm: float = None) -> str:
    """
    Adds a picture from a local path to the document.

    Args:
        filename (str): Path to the Word document.
        picture_path (str): File path of the image to be added.
        width_cm (float, optional): Width of the picture in centimeters.

    Returns:
        str: A JSON string confirming the operation.

    Example:
        add_picture(filename='report.docx', picture_path='logo.png', width_cm=5)
    """
    try:
        abs_doc_path = _get_abs_path(filename)
        abs_pic_path = _get_abs_path(picture_path)
        document = docx.Document(abs_doc_path)
        width = Cm(width_cm) if width_cm else None
        document.add_picture(abs_pic_path, width=width)
        document.save(abs_doc_path)
        return json.dumps({'status': 'success'})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def add_page_break(filename: str) -> str:
    """
    Adds a manual page break to the end of the document.

    Args:
        filename (str): The path to the Word document.

    Returns:
        str: A JSON string confirming the operation.

    Example:
        add_page_break(filename='report.docx')
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        document.add_page_break()
        document.save(abs_path)
        return json.dumps({'status': 'success'})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def add_header(filename: str, text: str) -> str:
    """
    Adds or replaces the document's primary header.

    Args:
        filename (str): The path to the Word document.
        text (str): The text to place in the header.

    Returns:
        str: A JSON string confirming the operation.

    Example:
        add_header(filename='report.docx', text='Company Confidential')
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        section = document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.text = text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.save(abs_path)
        return json.dumps({'status': 'success'})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def add_footer(filename: str, text: str) -> str:
    """
    Adds or replaces the document's primary footer.

    Args:
        filename (str): The path to the Word document.
        text (str): The text to place in the footer.

    Returns:
        str: A JSON string confirming the operation.

    Example:
        add_footer(filename='report.docx', text='Page 1')
    """
    try:
        abs_path = _get_abs_path(filename)
        document = docx.Document(abs_path)
        section = document.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.text = text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.save(abs_path)
        return json.dumps({'status': 'success'})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})

@mcp.tool()
def convert_to_pdf(filename: str, output_filename: str = None) -> str:
    """
    Converts a Word document to PDF format.

    Args:
        filename (str): Path to the source Word document.
        output_filename (str, optional): Path for the output PDF file.
                                        Defaults to the same name with a .pdf extension.

    Returns:
        str: A JSON string confirming the conversion and providing the PDF path.

    Example:
        convert_to_pdf(filename='report.docx', output_filename='report.pdf')
    """
    try:
        abs_input_path = _get_abs_path(filename)
        if output_filename:
            abs_output_path = _get_abs_path(output_filename)
        else:
            abs_output_path = os.path.splitext(abs_input_path)[0] + ".pdf"

        docx2pdf.convert(abs_input_path, abs_output_path)
        return json.dumps({'status': 'success', 'pdf_path': abs_output_path})
    except Exception as e:
        return json.dumps({'status': 'error', 'message': str(e)})


if __name__ == "__main__":
    # Ensure UTF-8 encoding for stdout
    if sys.stdout.encoding != 'utf-8':
        sys.stdout.reconfigure(encoding='utf-8')
    mcp.run()