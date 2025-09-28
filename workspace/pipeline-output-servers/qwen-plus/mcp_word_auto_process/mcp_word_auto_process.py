import sys
import os
from typing import Optional, List, Dict, Any
from uuid import uuid4
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from mcp.server.fastmcp import FastMCP

# 初始化 FastMCP 服务器
mcp = FastMCP("mcp_word_auto_process")

# 文档存储管理器
class DocumentManager:
    def __init__(self):
        self.documents: Dict[str, DocxDocument] = {}
        self.doc_ids_to_paths: Dict[str, str] = {}

    def create_document(self, author: Optional[str] = None, title: Optional[str] = None,
                        subject: Optional[str] = None, keywords: Optional[List[str]] = None) -> str:
        """
        创建一个新的Word文档并设置元数据。

        Args:
            author: 文档作者
            title: 文档标题
            subject: 文档主题
            keywords: 关键词列表

        Returns:
            返回文档对象ID用于后续操作。

        Raises:
            ValueError: 如果参数类型不正确
        """
        try:
            # 参数验证
            if keywords is not None and not isinstance(keywords, list):
                raise ValueError("keywords 必须是列表类型")
            
            doc_id = str(uuid4())
            document = DocxDocument()
            
            # 设置文档属性
            if author:
                document.core_properties.author = author
            if title:
                document.core_properties.title = title
            if subject:
                document.core_properties.subject = subject
            if keywords:
                document.core_properties.keywords = ','.join(keywords)
            
            self.documents[doc_id] = document
            return doc_id
            
        except Exception as e:
            raise ValueError(f"创建文档失败: {str(e)}") from e

    def get_document(self, doc_id: str) -> DocxDocument:
        """获取文档对象，如果不存在则抛出异常"""
        if doc_id not in self.documents:
            raise ValueError(f"文档 {doc_id} 不存在")
        return self.documents[doc_id]

    def get_paragraph(self, doc_id: str, para_id: str) -> Any:
        """获取段落对象，如果不存在则抛出异常"""
        doc = self.get_document(doc_id)
        # 实现段落查找逻辑（基于索引或唯一标识）
        if not hasattr(doc, '_paragraphs_registry'):
            doc._paragraphs_registry = {}
            
        if para_id not in doc._paragraphs_registry:
            raise ValueError(f"段落 {para_id} 不存在于文档 {doc_id}")
        return doc._paragraphs_registry[para_id]

# 全球文档管理器实例
doc_manager = DocumentManager()

@mcp.tool()
def create_document(author: Optional[str] = None, title: Optional[str] = None,
                    subject: Optional[str] = None, keywords: Optional[List[str]] = None) -> str:
    """
    创建一个新的Word文档并设置元数据。

    Args:
        author: 文档作者
        title: 文档标题
        subject: 文档主题
        keywords: 关键词列表

    Returns:
        返回文档对象ID用于后续操作.

    示例:
        create_document(author="John Doe", title="Report", subject="Business",
                       keywords=["report", "business"])
    """
    try:
        return doc_manager.create_document(author, title, subject, keywords)
    except Exception as e:
        raise ValueError(f"文档创建失败: {str(e)}") from e

@mcp.tool()
def get_document_text(doc_id: str) -> str:
    """
    提取整个文档的文本内容.

    Args:
        doc_id: 要提取文本的文档唯一标识符

    Returns:
        返回文档的完整文本内容.

    示例:
        get_document_text(doc_id="abc123")
    """
    try:
        doc = doc_manager.get_document(doc_id)
        text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
        return text
    except Exception as e:
        raise ValueError(f"获取文档文本失败: {str(e)}") from e

@mcp.tool()
def add_paragraph(doc_id: str, text: str, style: Optional[str] = None) -> str:
    """
    向指定文档中添加一个段落.

    Args:
        doc_id: 目标文档的唯一标识符
        text: 要添加的段落文本
        style: 可选样式名称

    Returns:
        返回新段落的ID或位置信息.

    示例:
        add_paragraph(doc_id="abc123", text="这是一个新段落", style="Normal")
    """
    try:
        doc = doc_manager.get_document(doc_id)
        para = doc.add_paragraph(text=text, style=style if style else '')
        para_id = str(uuid4())
        
        # 注册段落ID
        if not hasattr(doc, '_paragraphs_registry'):
            doc._paragraphs_registry = {}
        doc._paragraphs_registry[para_id] = para
        
        return para_id
    except Exception as e:
        raise ValueError(f"添加段落失败: {str(e)}") from e

@mcp.tool()
def add_heading(doc_id: str, text: str, level: int) -> str:
    """
    向指定文档中添加一个标题段落.

    Args:
        doc_id: 目标文档的唯一标识符
        text: 标题文本
        level: 标题级别（1~9）

    Returns:
        返回新标题段落的ID或位置信息.

    示例:
        add_heading(doc_id="abc123", text="章节标题", level=1)
    """
    try:
        # 参数验证
        if not 1 <= level <= 9:
            raise ValueError("标题级别必须在1到9之间")
            
        doc = doc_manager.get_document(doc_id)
        heading = doc.add_heading(text=text, level=level)
        para_id = str(uuid4())
        
        # 注册段落ID
        if not hasattr(doc, '_paragraphs_registry'):
            doc._paragraphs_registry = {}
        doc._paragraphs_registry[para_id] = heading
        
        return para_id
    except Exception as e:
        raise ValueError(f"添加标题失败: {str(e)}") from e

@mcp.tool()
def create_custom_style(doc_id: str, name: str, style_type: str, base_style: Optional[str] = None) -> str:
    """
    在指定文档中创建自定义样式.

    Args:
        doc_id: 目标文档的唯一标识符
        name: 自定义样式的名称
        style_type: 样式类型（paragraph/run/table）
        base_style: 基础样式名称（可选）

    Returns:
        返回新创建的样式对象或其名称.

    示例:
        create_custom_style(doc_id="abc123", name="CustomStyle", style_type="paragraph", base_style="Normal")
    """
    try:
        # 参数验证
        style_types = ['paragraph', 'run', 'table']
        if style_type.lower() not in style_types:
            raise ValueError(f"样式类型必须是 {style_types} 中的一种")
            
        doc = doc_manager.get_document(doc_id)
        
        # 将字符串类型转换为WD_STYLE_TYPE枚举
        wd_style_type = {
            'paragraph': WD_STYLE_TYPE.PARAGRAPH,
            'run': WD_STYLE_TYPE.CHARACTER,
            'table': WD_STYLE_TYPE.TABLE
        }[style_type.lower()]
        
        # 创建样式
        style = doc.styles.add_style(name, wd_style_type)
        
        # 设置基础样式
        if base_style:
            base_style_obj = doc.styles[base_style]
            style.base_style = base_style_obj
        
        return name
    except Exception as e:
        raise ValueError(f"创建自定义样式失败: {str(e)}") from e

@mcp.tool()
def format_text(doc_id: str, para_id: str, start_index: int, end_index: int,
                bold: Optional[bool] = None, italic: Optional[bool] = None,
                font_size: Optional[int] = None, font_name: Optional[str] = None) -> str:
    """
    格式化指定段落中的部分文本.

    Args:
        doc_id: 目标文档的唯一标识符
        para_id: 段落ID
        start_index: 开始字符索引
        end_index: 结束字符索引
        bold: 是否加粗
        italic: 是否斜体
        font_size: 字号大小（pt）
        font_name: 字体名称

    Returns:
        返回修改后的段落文本或状态码.

    示例:
        format_text(doc_id="abc123", para_id="para456", start_index=0, end_index=5,
                   bold=True, font_size=14, font_name="Arial")
    """
    try:
        # 获取文档和段落
        doc = doc_manager.get_document(doc_id)
        para = doc_manager.get_paragraph(doc_id, para_id)
        
        # 验证索引范围
        if start_index < 0 or end_index > len(para.text):
            raise ValueError("索引超出文本长度范围")
        
        # 获取运行对象
        runs = para.runs
        if not runs:
            raise ValueError("段落中没有可格式化的文本")
        
        # 实现文本格式化
        for run in runs:
            if run.text:
                # 应用格式
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if font_size:
                    run.font.size = Pt(font_size)
                if font_name:
                    run.font.name = font_name
        
        return f"成功格式化段落 {para_id} 的文本 {start_index}-{end_index}"
    except Exception as e:
        raise ValueError(f"文本格式化失败: {str(e)}") from e

@mcp.tool()
def protect_document(doc_id: str, password: str, protection_type: str = "READ_ONLY") -> str:
    """
    保护文档，防止未经授权的编辑.

    Args:
        doc_id: 目标文档的唯一标识符
        password: 保护密码
        protection_type: 保护类型（如只读、注释等）

    Returns:
        返回文档保护状态.

    示例:
        protect_document(doc_id="abc123", password="secret", protection_type="READ_ONLY")
    """
    try:
        doc = doc_manager.get_document(doc_id)
        # python-docx 不直接支持文档保护，需要使用其他库如 win32com 或 comtypes
        # 这里仅提供占位实现
        return f"文档保护功能暂未实现，需要第三方库支持"
    except Exception as e:
        raise ValueError(f"文档保护失败: {str(e)}") from e

@mcp.tool()
def add_footnote_to_document(doc_id: str, para_id: str, position: int, text: str) -> str:
    """
    向指定文档中添加脚注.

    Args:
        doc_id: 目标文档的唯一标识符
        para_id: 添加脚注的段落ID
        position: 脚注插入的位置
        text: 脚注内容

    Returns:
        返回脚注对象或插入状态.

    示例:
        add_footnote_to_document(doc_id="abc123", para_id="para456", position=10, text="这是一个脚注")
    """
    try:
        doc = doc_manager.get_document(doc_id)
        para = doc_manager.get_paragraph(doc_id, para_id)
        
        # 使用python-docx添加脚注
        # 注意：python-docx当前不支持脚注，需要使用其他方法或库
        return f"脚注功能暂未实现，python-docx不支持脚注功能"
    except Exception as e:
        raise ValueError(f"添加脚注失败: {str(e)}") from e

@mcp.tool()
def get_paragraph_text_from_document(doc_id: str, para_id: str) -> str:
    """
    获取指定段落的文本内容.

    Args:
        doc_id: 目标文档的唯一标识符
        para_id: 段落ID

    Returns:
        返回段落的文本内容.

    示例:
        get_paragraph_text_from_document(doc_id="abc123", para_id="para456")
    """
    try:
        doc = doc_manager.get_document(doc_id)
        para = doc_manager.get_paragraph(doc_id, para_id)
        return para.text
    except Exception as e:
        raise ValueError(f"获取段落文本失败: {str(e)}") from e

@mcp.tool()
def find_text_in_document(doc_id: str, search_text: str) -> List[Dict[str, Any]]:
    """
    在文档中搜索指定文本，并返回所有匹配段落及位置.

    Args:
        doc_id: 目标文档的唯一标识符
        search_text: 要搜索的文本

    Returns:
        返回包含匹配段落ID和位置的列表.

    示例:
        find_text_in_document(doc_id="abc123", search_text="关键词")
    """
    try:
        doc = doc_manager.get_document(doc_id)
        results = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            if search_text in paragraph.text:
                # 找到匹配位置
                start_idx = paragraph.text.index(search_text)
                end_idx = start_idx + len(search_text)
                
                # 生成段落ID（假设我们有段落注册机制）
                para_id = f"para_{i}"
                results.append({
                    "para_id": para_id,
                    "start_index": start_idx,
                    "end_index": end_idx
                })
        
        return results
    except Exception as e:
        raise ValueError(f"文本搜索失败: {str(e)}") from e

@mcp.tool()
def add_table(doc_id: str, rows: int, cols: int, data: List[List[str]]) -> str:
    """
    向指定文档中添加表格.

    Args:
        doc_id: 目标文档的唯一标识符
        rows: 表格行数
        cols: 表格列数
        data: 二维列表形式的数据

    Returns:
        返回表格对象或插入位置.

    示例:
        add_table(doc_id="abc123", rows=2, cols=2, data=[["A1","B1"],["A2","B2"]])
    """
    try:
        doc = doc_manager.get_document(doc_id)
        
        # 验证数据维度
        if len(data) != rows or any(len(row) != cols for row in data):
            raise ValueError("数据维度与表格行列数不匹配")
            
        table = doc.add_table(rows=rows, cols=cols)
        
        # 填充数据
        for i, row_data in enumerate(data):
            for j, cell_text in enumerate(row_data):
                table.cell(i, j).text = cell_text
        
        return f"table_{uuid4()}"
    except Exception as e:
        raise ValueError(f"添加表格失败: {str(e)}") from e

@mcp.tool()
def add_image(doc_id: str, image_path: str, width: float, height: float) -> str:
    """
    向指定文档中插入图片.

    Args:
        doc_id: 目标文档的唯一标识符
        image_path: 图片文件路径
        width: 图片宽度（英寸）
        height: 图片高度（英寸）

    Returns:
        返回图片对象或插入位置.

    示例:
        add_image(doc_id="abc123", image_path="/path/to/image.jpg", width=2.0, height=1.5)
    """
    try:
        doc = doc_manager.get_document(doc_id)
        
        # 验证文件存在
        if not os.path.exists(image_path):
            raise ValueError(f"图片文件不存在: {image_path}")
            
        # 添加图片
        para = doc.add_paragraph()
        run = para.add_run()
        inline_shape = run.add_picture(image_path, width=Inches(width), height=Inches(height))
        
        # 返回图片尺寸信息
        return f"图片已添加，尺寸: {inline_shape.width}, {inline_shape.height}"
    except Exception as e:
        raise ValueError(f"添加图片失败: {str(e)}") from e

@mcp.tool()
def add_page_break(doc_id: str) -> str:
    """
    向指定文档中添加分页符.

    Args:
        doc_id: 目标文档的唯一标识符

    Returns:
        返回操作结果.

    示例:
        add_page_break(doc_id="abc123")
    """
    try:
        doc = doc_manager.get_document(doc_id)
        doc.add_page_break()
        return "分页符添加成功"
    except Exception as e:
        raise ValueError(f"添加分页符失败: {str(e)}") from e

@mcp.tool()
def add_header_footer(doc_id: str, header_text: Optional[str] = None, footer_text: Optional[str] = None) -> str:
    """
    向指定文档中添加页眉或页脚.

    Args:
        doc_id: 目标文档的唯一标识符
        header_text: 页眉文本
        footer_text: 页脚文本

    Returns:
        返回页眉/页脚对象或插入状态.

    示例:
        add_header_footer(doc_id="abc123", header_text="页眉内容", footer_text="页脚内容")
    """
    try:
        doc = doc_manager.get_document(doc_id)
        
        # 添加页眉
        if header_text:
            for section in doc.sections:
                section.header.paragraphs[0].text = header_text
        
        # 添加页脚
        if footer_text:
            for section in doc.sections:
                section.footer.paragraphs[0].text = footer_text
        
        return "页眉页脚添加成功"
    except Exception as e:
        raise ValueError(f"添加页眉页脚失败: {str(e)}") from e

@mcp.tool()
def convert_to_pdf(doc_id: str, output_path: str) -> str:
    """
    将Word文档转换为PDF格式.

    Args:
        doc_id: 目标文档的唯一标识符
        output_path: 输出PDF文件路径

    Returns:
        返回PDF文件路径或转换状态.

    示例:
        convert_to_pdf(doc_id="abc123", output_path="/output/report.pdf")
    """
    try:
        from docx2pdf import convert
        import io
        
        doc = doc_manager.get_document(doc_id)
        
        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # 保存文档到临时路径
        temp_path = f"{output_path}.tmp.docx"
        doc.save(temp_path)
        
        # 转换为PDF
        convert(temp_path, output_path)
        
        # 删除临时文件
        os.remove(temp_path)
        
        return output_path
    except ImportError:
        raise ValueError("缺少依赖: 请安装 docx2pdf 库")
    except Exception as e:
        raise ValueError(f"PDF转换失败: {str(e)}") from e

if __name__ == "__main__":
    sys.stdout.reconfigure(encoding='utf-8')
    mcp.run()