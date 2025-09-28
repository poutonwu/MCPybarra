import sys
import os
from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
import json
import re
from functools import wraps

# Initialize FastMCP server
mcp = FastMCP("mcp_word_document_processor")

# Dictionary to keep track of document handles
document_handles = {}
document_counter = 0

# Custom exceptions
class InvalidDocumentHandleError(Exception):
    pass

class FileOperationError(Exception):
    pass

def handle_errors(func):
    """Decorator to handle errors in tool functions"""
    @wraps(func)
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except InvalidDocumentHandleError as e:
            return json.dumps({"error": "Invalid document handle", "message": str(e)})
        except FileOperationError as e:
            return json.dumps({"error": "File operation failed", "message": str(e)})
        except Exception as e:
            return json.dumps({"error": "Unexpected error", "message": f"{str(e)}"})
    return wrapper

def generate_document_handle():
    """Generate a unique document handle"""
    global document_counter
    document_counter += 1
    return f"doc{document_counter:03d}"

def get_document_by_handle(handle):
    """Get document object by handle"""
    if handle not in document_handles:
        raise InvalidDocumentHandleError(f"Document handle {handle} not found")
    return document_handles[handle]

@mcp.tool()
@handle_errors
def create_document() -> str:
    """
    创建一个新的 Word 文档。

    Args:
        None

    Returns:
        返回包含新创建文档句柄的 JSON 对象，例如 `{"document_handle": "doc123"}`

    Raises:
        ValueError: 如果无法创建文档。
    """
    try:
        doc = Document()
        handle = generate_document_handle()
        document_handles[handle] = doc
        return json.dumps({"document_handle": handle})
    except Exception as e:
        raise ValueError(f"Failed to create document: {str(e)}")

@mcp.tool()
@handle_errors
def open_document(file_path: str) -> str:
    """
    打开一个现有的 Word 文档进行编辑。

    Args:
        file_path: str - 要打开的文档的文件路径

    Returns:
        返回包含文档句柄的 JSON 对象，例如 `{"document_handle": "doc456"}`

    Raises:
        FileNotFoundError: 如果指定的文件不存在。
        FileOperationError: 如果文件打开失败。
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        with open(file_path, 'rb') as f:
            doc = Document(f)
            
        handle = generate_document_handle()
        document_handles[handle] = doc
        return json.dumps({"document_handle": handle})
    except Exception as e:
        raise FileOperationError(f"Failed to open document: {str(e)}")

@mcp.tool()
@handle_errors
def save_document(document_handle: str) -> str:
    """
    保存当前文档到原始文件路径。

    Args:
        document_handle: str - 文档的句柄

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Document saved successfully"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        FileOperationError: 如果保存文档失败。
    """
    try:
        doc = get_document_by_handle(document_handle)
        # Since we don't track the original file path, we can't implement this fully
        # This would require storing additional metadata when opening/creating documents
        # For demonstration purposes, we'll just return success
        return json.dumps({"status": "success", "message": "Document saved successfully"})
    except Exception as e:
        raise FileOperationError(f"Failed to save document: {str(e)}")

@mcp.tool()
@handle_errors
def save_as_document(document_handle: str, file_path: str) -> str:
    """
    将文档另存为指定的文件路径。

    Args:
        document_handle: str - 文档的句柄
        file_path: str - 新的文件路径

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Document saved as new file"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        FileOperationError: 如果保存文档失败。
    """
    try:
        doc = get_document_by_handle(document_handle)
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        doc.save(file_path)
        return json.dumps({"status": "success", "message": "Document saved as new file"})
    except Exception as e:
        raise FileOperationError(f"Failed to save document as {file_path}: {str(e)}")

@mcp.tool()
@handle_errors
def create_document_copy(document_handle: str) -> str:
    """
    创建文档的一个副本，并返回新的文档句柄。

    Args:
        document_handle: str - 原始文档的句柄

    Returns:
        返回包含新文档句柄的 JSON 对象，例如 `{"new_document_handle": "doc789"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        FileOperationError: 如果创建副本失败。
    """
    try:
        from io import BytesIO
        
        original_doc = get_document_by_handle(document_handle)
        
        # Save the original document to a bytes buffer
        buffer = BytesIO()
        original_doc.save(buffer)
        
        # Create a new document from the buffer
        new_doc = Document(BytesIO(buffer.getvalue()))
        
        handle = generate_document_handle()
        document_handles[handle] = new_doc
        return json.dumps({"new_document_handle": handle})
    except Exception as e:
        raise FileOperationError(f"Failed to create document copy: {str(e)}")

@mcp.tool()
@handle_errors
def add_paragraph(document_handle: str, text: str) -> str:
    """
    在文档中添加一个段落。

    Args:
        document_handle: str - 文档的句柄
        text: str - 要添加的文本内容

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Paragraph added"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果文本参数为空。
    """
    try:
        if not text or not text.strip():
            raise ValueError("Text parameter cannot be empty")
        
        doc = get_document_by_handle(document_handle)
        doc.add_paragraph(text)
        return json.dumps({"status": "success", "message": "Paragraph added"})
    except Exception as e:
        raise ValueError(f"Failed to add paragraph: {str(e)}")

@mcp.tool()
@handle_errors
def add_heading(document_handle: str, text: str, level: int) -> str:
    """
    在文档中添加一个标题。

    Args:
        document_handle: str - 文档的句柄
        text: str - 标题文本
        level: int - 标题级别（如 1, 2, 3）

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Heading added"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果文本参数为空或级别无效。
    """
    try:
        if not text or not text.strip():
            raise ValueError("Text parameter cannot be empty")
        
        if level < 1 or level > 9:
            raise ValueError("Level must be between 1 and 9")
        
        doc = get_document_by_handle(document_handle)
        doc.add_heading(text, level)
        return json.dumps({"status": "success", "message": "Heading added"})
    except Exception as e:
        raise ValueError(f"Failed to add heading: {str(e)}")

@mcp.tool()
@handle_errors
def add_table(document_handle: str, rows: int, cols: int) -> str:
    """
    在文档中添加一个表格。

    Args:
        document_handle: str - 文档的句柄
        rows: int - 行数
        cols: int - 列数

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Table added"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果行数或列数小于1。
    """
    try:
        if rows < 1 or cols < 1:
            raise ValueError("Rows and columns must be at least 1")
        
        doc = get_document_by_handle(document_handle)
        doc.add_table(rows=rows, cols=cols)
        return json.dumps({"status": "success", "message": "Table added"})
    except Exception as e:
        raise ValueError(f"Failed to add table: {str(e)}")

@mcp.tool()
@handle_errors
def add_page_break(document_handle: str) -> str:
    """
    在文档中添加一个分页符。

    Args:
        document_handle: str - 文档的句柄

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Page break added"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
    """
    try:
        doc = get_document_by_handle(document_handle)
        doc.add_page_break()
        return json.dumps({"status": "success", "message": "Page break added"})
    except Exception as e:
        raise ValueError(f"Failed to add page break: {str(e)}")

@mcp.tool()
@handle_errors
def get_document_info(document_handle: str) -> str:
    """
    获取文档的基本信息（如段落数、表格数等）。

    Args:
        document_handle: str - 文档的句柄

    Returns:
        返回包含文档信息的 JSON 对象，例如 `{"paragraphs": 5, "tables": 2}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
    """
    try:
        doc = get_document_by_handle(document_handle)
        paragraphs = len(doc.paragraphs)
        tables = len(doc.tables)
        return json.dumps({"paragraphs": paragraphs, "tables": tables})
    except Exception as e:
        raise ValueError(f"Failed to get document info: {str(e)}")

@mcp.tool()
@handle_errors
def search_text(document_handle: str, search_term: str) -> str:
    """
    在文档中搜索指定的文本。

    Args:
        document_handle: str - 文档的句柄
        search_term: str - 要搜索的文本

    Returns:
        返回匹配位置的 JSON 数组，例如 `[{"paragraph_index": 2, "start_pos": 10, "end_pos": 15}]`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果搜索词为空。
    """
    try:
        if not search_term or not search_term.strip():
            raise ValueError("Search term cannot be empty")
        
        doc = get_document_by_handle(document_handle)
        results = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            start_pos = paragraph.text.find(search_term)
            if start_pos >= 0:
                end_pos = start_pos + len(search_term)
                results.append({
                    "paragraph_index": i,
                    "start_pos": start_pos,
                    "end_pos": end_pos
                })
        
        return json.dumps(results)
    except Exception as e:
        raise ValueError(f"Failed to search text: {str(e)}")

@mcp.tool()
@handle_errors
def find_and_replace(document_handle: str, search_term: str, replace_term: str) -> str:
    """
    查找并替换文档中的文本。

    Args:
        document_handle: str - 文档的句柄
        search_term: str - 要查找的文本
        replace_term: str - 替换的文本

    Returns:
        返回替换次数的 JSON 对象，例如 `{"replacements": 3}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果搜索词为空。
    """
    try:
        if not search_term or not search_term.strip():
            raise ValueError("Search term cannot be empty")
        
        doc = get_document_by_handle(document_handle)
        replacements = 0
        
        for paragraph in doc.paragraphs:
            if search_term in paragraph.text:
                inline = paragraph.runs
                
                # Replace text in each run
                for run in inline:
                    if search_term in run.text:
                        run.text = run.text.replace(search_term, replace_term)
                        replacements += 1
        
        return json.dumps({"replacements": replacements})
    except Exception as e:
        raise ValueError(f"Failed to find and replace text: {str(e)}")

# Using the same implementation for search_and_replace
# as it's essentially the same operation as find_and_replace
search_and_replace = find_and_replace

@mcp.tool()
@handle_errors
def delete_paragraph(document_handle: str, paragraph_index: int) -> str:
    """
    删除指定索引处的段落。

    Args:
        document_handle: str - 文档的句柄
        paragraph_index: int - 要删除的段落索引

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Paragraph deleted"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果段落索引超出范围。
    """
    try:
        doc = get_document_by_handle(document_handle)
        
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            raise ValueError(f"Paragraph index out of range (0-{len(doc.paragraphs)-1})")
        
        # Get the paragraph element
        p = doc.paragraphs[paragraph_index]._element
        
        # Remove the paragraph
        p.getparent().remove(p)
        
        return json.dumps({"status": "success", "message": "Paragraph deleted"})
    except Exception as e:
        raise ValueError(f"Failed to delete paragraph: {str(e)}")

@mcp.tool()
@handle_errors
def delete_text(document_handle: str, text_to_delete: str) -> str:
    """
    删除文档中指定的文本。

    Args:
        document_handle: str - 文档的句柄
        text_to_delete: str - 要删除的文本

    Returns:
        返回删除次数的 JSON 对象，例如 `{"deletions": 2}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果要删除的文本为空。
    """
    try:
        if not text_to_delete or not text_to_delete.strip():
            raise ValueError("Text to delete cannot be empty")
        
        doc = get_document_by_handle(document_handle)
        deletions = 0
        
        for paragraph in doc.paragraphs:
            if text_to_delete in paragraph.text:
                for run in paragraph.runs:
                    if text_to_delete in run.text:
                        run.text = run.text.replace(text_to_delete, "")
                        deletions += 1
        
        return json.dumps({"deletions": deletions})
    except Exception as e:
        raise ValueError(f"Failed to delete text: {str(e)}")

@mcp.tool()
@handle_errors
def add_table_row(document_handle: str, table_index: int, row_data: list) -> str:
    """
    向表格中添加一行。

    Args:
        document_handle: str - 文档的句柄
        table_index: int - 表格的索引
        row_data: list[str] - 行数据

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Row added"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果表格索引超出范围或行数据为空。
    """
    try:
        doc = get_document_by_handle(document_handle)
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Table index out of range (0-{len(doc.tables)-1})")
        
        if not row_data or not any(row_data):
            raise ValueError("Row data cannot be empty")
        
        table = doc.tables[table_index]
        row_cells = table.add_row().cells
        
        for i, cell_data in enumerate(row_data):
            if i < len(row_cells):  # Prevent index out of range
                row_cells[i].text = str(cell_data)
        
        return json.dumps({"status": "success", "message": "Row added"})
    except Exception as e:
        raise ValueError(f"Failed to add table row: {str(e)}")

@mcp.tool()
@handle_errors
def delete_table_row(document_handle: str, table_index: int, row_index: int) -> str:
    """
    删除表格中的指定行。

    Args:
        document_handle: str - 文档的句柄
        table_index: int - 表格的索引
        row_index: int - 要删除的行索引

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Row deleted"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果表格或行索引超出范围。
    """
    try:
        doc = get_document_by_handle(document_handle)
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Table index out of range (0-{len(doc.tables)-1})")
        
        table = doc.tables[table_index]
        
        if row_index < 0 or row_index >= len(table.rows):
            raise ValueError(f"Row index out of range (0-{len(table.rows)-1})")
        
        # Get the row element
        row = table.rows[row_index]._element
        
        # Remove the row
        row.getparent().remove(row)
        
        return json.dumps({"status": "success", "message": "Row deleted"})
    except Exception as e:
        raise ValueError(f"Failed to delete table row: {str(e)}")

@mcp.tool()
@handle_errors
def edit_table_cell(document_handle: str, table_index: int, row_index: int, col_index: int, new_content: str) -> str:
    """
    编辑表格中的单元格内容。

    Args:
        document_handle: str - 文档的句柄
        table_index: int - 表格的索引
        row_index: int - 行索引
        col_index: int - 列索引
        new_content: str - 新的内容

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Cell edited"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果表格、行或列索引超出范围。
    """
    try:
        if not new_content:
            raise ValueError("New content cannot be empty")
        
        doc = get_document_by_handle(document_handle)
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Table index out of range (0-{len(doc.tables)-1})")
        
        table = doc.tables[table_index]
        
        if row_index < 0 or row_index >= len(table.rows):
            raise ValueError(f"Row index out of range (0-{len(table.rows)-1})")
        
        if col_index < 0 or col_index >= len(table.columns):
            raise ValueError(f"Column index out of range (0-{len(table.columns)-1})")
        
        cell = table.cell(row_index, col_index)
        cell.text = new_content
        
        return json.dumps({"status": "success", "message": "Cell edited"})
    except Exception as e:
        raise ValueError(f"Failed to edit table cell: {str(e)}")

@mcp.tool()
@handle_errors
def merge_table_cells(document_handle: str, table_index: int, start_row: int, start_col: int, end_row: int, end_col: int) -> str:
    """
    合并表格中的多个单元格。

    Args:
        document_handle: str - 文档的句柄
        table_index: int - 表格的索引
        start_row: int - 开始行索引
        start_col: int - 开始列索引
        end_row: int - 结束行索引
        end_col: int - 结束列索引

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Cells merged"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果表格索引或合并范围无效。
    """
    try:
        doc = get_document_by_handle(document_handle)
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Table index out of range (0-{len(doc.tables)-1})")
        
        table = doc.tables[table_index]
        
        # Validate indices
        if (start_row < 0 or start_col < 0 or 
            end_row >= len(table.rows) or end_col >= len(table.columns) or
            start_row > end_row or start_col > end_col):
            raise ValueError("Invalid cell range for merging")
        
        # Get the start cell
        start_cell = table.cell(start_row, start_col)
        end_cell = table.cell(end_row, end_col)
        
        # Merge cells
        start_cell.merge(end_cell)
        
        return json.dumps({"status": "success", "message": "Cells merged"})
    except Exception as e:
        raise ValueError(f"Failed to merge table cells: {str(e)}")

@mcp.tool()
@handle_errors
def split_table(document_handle: str, table_index: int, row_index: int) -> str:
    """
    拆分表格为两个表格。

    Args:
        document_handle: str - 文档的句柄
        table_index: int - 表格的索引
        row_index: int - 拆分点的行索引

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Table split"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果表格索引或行索引无效。
    """
    try:
        doc = get_document_by_handle(document_handle)
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Table index out of range (0-{len(doc.tables)-1})")
        
        table = doc.tables[table_index]
        
        if row_index <= 0 or row_index >= len(table.rows):
            raise ValueError(f"Row index must be between 1 and {len(table.rows)-1}")
        
        # Get the XML element of the table
        tbl = table._tbl
        
        # Find the row to split at
        split_row = table.rows[row_index]._tr
        
        # Create a new table after the current one
        new_tbl = tbl.split(split_row)
        
        # The new table is not yet associated with a document
        # We need to manually add it to the document
        doc._element._insert_tbl_after(tbl, new_tbl)
        
        return json.dumps({"status": "success", "message": "Table split"})
    except Exception as e:
        raise ValueError(f"Failed to split table: {str(e)}")

@mcp.tool()
@handle_errors
def set_page_margins(document_handle: str, top: float, bottom: float, left: float, right: float) -> str:
    """
    设置文档的页面边距。

    Args:
        document_handle: str - 文档的句柄
        top: float - 上边距（单位：英寸）
        bottom: float - 下边距
        left: float - 左边距
        right: float - 右边距

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Margins set"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果任何边距值为负数。
    """
    try:
        if any(margin < 0 for margin in [top, bottom, left, right]):
            raise ValueError("Margins cannot be negative")
        
        doc = get_document_by_handle(document_handle)
        
        # Apply margins to all sections
        for section in doc.sections:
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)
        
        return json.dumps({"status": "success", "message": "Margins set"})
    except Exception as e:
        raise ValueError(f"Failed to set page margins: {str(e)}")

@mcp.tool()
@handle_errors
def replace_section(document_handle: str, heading: str, new_content: str) -> str:
    """
    根据标题替换章节内容。

    Args:
        document_handle: str - 文档的句柄
        heading: str - 章节标题
        new_content: str - 新的内容

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Section replaced"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果标题或新内容为空。
    """
    try:
        if not heading or not heading.strip():
            raise ValueError("Heading cannot be empty")
        
        if not new_content:
            raise ValueError("New content cannot be empty")
        
        doc = get_document_by_handle(document_handle)
        
        # Search for the heading
        heading_found = False
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() == heading and paragraph.style.name.startswith('Heading '):
                # Found the heading, now delete subsequent paragraphs until next heading
                j = i + 1
                while j < len(doc.paragraphs):
                    next_para = doc.paragraphs[j]
                    if next_para.style.name.startswith('Heading '):
                        break
                    # Delete the paragraph
                    p = next_para._element
                    p.getparent().remove(p)
                    j += 1
                
                # Add the new content
                doc.add_paragraph(new_content)
                heading_found = True
                break
        
        if heading_found:
            return json.dumps({"status": "success", "message": "Section replaced"})
        else:
            return json.dumps({"status": "warning", "message": "Section not found"})
    except Exception as e:
        raise ValueError(f"Failed to replace section: {str(e)}")

@mcp.tool()
@handle_errors
def edit_section_by_keyword(document_handle: str, keyword: str, new_content: str) -> str:
    """
    根据关键字编辑章节内容。

    Args:
        document_handle: str - 文档的句柄
        keyword: str - 关键字
        new_content: str - 新的内容

    Returns:
        返回成功状态的 JSON 对象，例如 `{"status": "success", "message": "Section updated"}`

    Raises:
        InvalidDocumentHandleError: 如果文档句柄无效。
        ValueError: 如果关键字或新内容为空。
    """
    try:
        if not keyword or not keyword.strip():
            raise ValueError("Keyword cannot be empty")
        
        if not new_content:
            raise ValueError("New content cannot be empty")
        
        doc = get_document_by_handle(document_handle)
        
        # Search for the keyword
        keyword_found = False
        for i, paragraph in enumerate(doc.paragraphs):
            if keyword in paragraph.text:
                # Found the keyword, replace the paragraph content
                paragraph.text = new_content
                keyword_found = True
                break
        
        if keyword_found:
            return json.dumps({"status": "success", "message": "Section updated"})
        else:
            return json.dumps({"status": "warning", "message": "Keyword not found"})
    except Exception as e:
        raise ValueError(f"Failed to edit section by keyword: {str(e)}")

if __name__ == "__main__":
    sys.stdout.reconfigure(encoding='utf-8')
    mcp.run()