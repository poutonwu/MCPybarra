import os
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from mcp.server.fastmcp import FastMCP

# Initialize FastMCP server
mcp = FastMCP("mcp_word_document_automation")

class WordDocumentManager:
    def __init__(self):
        self.document = None

    def create_document(self, file_name: str) -> str:
        """Creates a new Word document.

        Args:
            file_name (str): The name of the file to be created.

        Returns:
            A confirmation message indicating the document was created successfully.
        """
        try:
            if not file_name.endswith('.docx'):
                file_name += '.docx'
            self.document = Document()
            return f"Document '{file_name}' created successfully."
        except Exception as e:
            return f"Error creating document: {e}"

    def open_document(self, file_path: str) -> str:
        """Opens an existing Word document.

        Args:
            file_path (str): The path to the document file.

        Returns:
            A confirmation message indicating the document was opened successfully.
        """
        try:
            if not os.path.exists(file_path):
                return f"File '{file_path}' does not exist."
            self.document = Document(file_path)
            return f"Document '{file_path}' opened successfully."
        except Exception as e:
            return f"Error opening document: {e}"

    def save_document(self, file_path: str) -> str:
        """Saves the currently open Word document.

        Args:
            file_path (str): The path where the document should be saved.

        Returns:
            A confirmation message indicating the document was saved successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            self.document.save(file_path)
            return f"Document saved successfully at '{file_path}'."
        except Exception as e:
            return f"Error saving document: {e}"

    def save_as_document(self, new_file_path: str) -> str:
        """Saves the currently open Word document with a new name or in a different location.

        Args:
            new_file_path (str): The new file path for saving the document.

        Returns:
            A confirmation message indicating the document was saved successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            if not new_file_path.endswith('.docx'):
                new_file_path += '.docx'
            self.document.save(new_file_path)
            return f"Document saved as '{new_file_path}' successfully."
        except Exception as e:
            return f"Error saving document: {e}"

    def create_document_copy(self, copy_file_path: str) -> str:
        """Creates a copy of the currently open Word document.

        Args:
            copy_file_path (str): The file path where the copy should be saved.

        Returns:
            A confirmation message indicating the document copy was created successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            if not copy_file_path.endswith('.docx'):
                copy_file_path += '.docx'
            self.document.save(copy_file_path)
            return f"Document copy created at '{copy_file_path}' successfully."
        except Exception as e:
            return f"Error creating document copy: {e}"

    def add_paragraph(self, text: str) -> str:
        """Adds a paragraph to the document.

        Args:
            text (str): The text content of the paragraph.

        Returns:
            A confirmation message indicating the paragraph was added successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            self.document.add_paragraph(text)
            return f"Paragraph added successfully."
        except Exception as e:
            return f"Error adding paragraph: {e}"

    def add_heading(self, text: str, level: int) -> str:
        """Adds a heading to the document.

        Args:
            text (str): The text content of the heading.
            level (int): The level of the heading (e.g., 1 for main heading, 2 for subheading).

        Returns:
            A confirmation message indicating the heading was added successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            if level < 1 or level > 9:
                return "Heading level must be between 1 and 9."
            self.document.add_heading(text, level=level)
            return f"Heading added successfully."
        except Exception as e:
            return f"Error adding heading: {e}"

    def add_table(self, rows: int, cols: int) -> str:
        """Adds a table to the document.

        Args:
            rows (int): The number of rows in the table.
            cols (int): The number of columns in the table.

        Returns:
            A confirmation message indicating the table was added successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            if rows <= 0 or cols <= 0:
                return "Rows and columns must be positive integers."
            self.document.add_table(rows=rows, cols=cols)
            return f"Table added successfully."
        except Exception as e:
            return f"Error adding table: {e}"

    def add_page_break(self) -> str:
        """Adds a page break to the document.

        Returns:
            A confirmation message indicating the page break was added successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            self.document.add_page_break()
            return "Page break added successfully."
        except Exception as e:
            return f"Error adding page break: {e}"

    def get_document_info(self) -> dict:
        """Retrieves metadata about the document, such as the number of pages, paragraphs, and tables.

        Returns:
            A dictionary containing metadata about the document.
        """
        try:
            if not self.document:
                return "No document is currently open."
            info = {
                "paragraphs": len(self.document.paragraphs),
                "tables": len(self.document.tables),
            }
            return info
        except Exception as e:
            return f"Error retrieving document info: {e}"

    def search_text(self, query: str) -> list:
        """Searches for specific text within the document.

        Args:
            query (str): The text to search for.

        Returns:
            A list of locations where the text was found.
        """
        try:
            if not self.document:
                return "No document is currently open."
            locations = []
            for i, paragraph in enumerate(self.document.paragraphs):
                if query in paragraph.text:
                    locations.append(f"Paragraph {i + 1}")
            for i, table in enumerate(self.document.tables):
                for j, row in enumerate(table.rows):
                    for k, cell in enumerate(row.cells):
                        if query in cell.text:
                            locations.append(f"Table {i + 1}, Row {j + 1}, Cell {k + 1}")
            return locations
        except Exception as e:
            return f"Error searching text: {e}"

    def find_and_replace(self, old_text: str, new_text: str) -> str:
        """Finds specific text in the document and replaces it with new text.

        Args:
            old_text (str): The text to find.
            new_text (str): The text to replace it with.

        Returns:
            A confirmation message indicating the replacement was successful.
        """
        try:
            if not self.document:
                return "No document is currently open."
            for paragraph in self.document.paragraphs:
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
            for table in self.document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)
            return "Text replaced successfully."
        except Exception as e:
            return f"Error replacing text: {e}"

    def delete_paragraph(self, paragraph_index: int) -> str:
        """Deletes a specified paragraph from the document.

        Args:
            paragraph_index (int): The index of the paragraph to delete.

        Returns:
            A confirmation message indicating the paragraph was deleted successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            if paragraph_index < 0 or paragraph_index >= len(self.document.paragraphs):
                return "Invalid paragraph index."
            p = self.document.paragraphs[paragraph_index]
            p._element.getparent().remove(p._element)
            return "Paragraph deleted successfully."
        except Exception as e:
            return f"Error deleting paragraph: {e}"

    def delete_text(self, text_to_delete: str) -> str:
        """Deletes all instances of specified text from the document.

        Args:
            text_to_delete (str): The text to delete.

        Returns:
            A confirmation message indicating the text was deleted successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            for paragraph in self.document.paragraphs:
                if text_to_delete in paragraph.text:
                    paragraph.text = paragraph.text.replace(text_to_delete, '')
            for table in self.document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if text_to_delete in cell.text:
                            cell.text = cell.text.replace(text_to_delete, '')
            return "Text deleted successfully."
        except Exception as e:
            return f"Error deleting text: {e}"

    def add_table_row(self, table_index: int) -> str:
        """Adds a new row to an existing table.

        Args:
            table_index (int): The index of the table.

        Returns:
            A confirmation message indicating the row was added successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            if table_index < 0 or table_index >= len(self.document.tables):
                return "Invalid table index."
            table = self.document.tables[table_index]
            table.add_row()
            return "Row added successfully."
        except Exception as e:
            return f"Error adding table row: {e}"

    def delete_table_row(self, table_index: int, row_index: int) -> str:
        """Deletes a row from an existing table.

        Args:
            table_index (int): The index of the table.
            row_index (int): The index of the row to delete.

        Returns:
            A confirmation message indicating the row was deleted successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            if table_index < 0 or table_index >= len(self.document.tables):
                return "Invalid table index."
            table = self.document.tables[table_index]
            if row_index < 0 or row_index >= len(table.rows):
                return "Invalid row index."
            row = table.rows[row_index]
            row._element.getparent().remove(row._element)
            return "Row deleted successfully."
        except Exception as e:
            return f"Error deleting table row: {e}"

    def edit_table_cell(self, table_index: int, row_index: int, col_index: int, new_content: str) -> str:
        """Edits the content of a specific cell in a table.

        Args:
            table_index (int): The index of the table.
            row_index (int): The index of the row.
            col_index (int): The index of the column.
            new_content (str): The new content for the cell.

        Returns:
            A confirmation message indicating the cell was edited successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            if table_index < 0 or table_index >= len(self.document.tables):
                return "Invalid table index."
            table = self.document.tables[table_index]
            if row_index < 0 or row_index >= len(table.rows):
                return "Invalid row index."
            if col_index < 0 or col_index >= len(table.columns):
                return "Invalid column index."
            cell = table.cell(row_index, col_index)
            cell.text = new_content
            return "Cell edited successfully."
        except Exception as e:
            return f"Error editing table cell: {e}"

    def merge_table_cells(self, table_index: int, start_row: int, start_col: int, end_row: int, end_col: int) -> str:
        """Merges cells in a table.

        Args:
            table_index (int): The index of the table.
            start_row (int): The starting row index for merging.
            start_col (int): The starting column index for merging.
            end_row (int): The ending row index for merging.
            end_col (int): The ending column index for merging.

        Returns:
            A confirmation message indicating the cells were merged successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            if table_index < 0 or table_index >= len(self.document.tables):
                return "Invalid table index."
            table = self.document.tables[table_index]
            start_cell = table.cell(start_row, start_col)
            end_cell = table.cell(end_row, end_col)
            start_cell.merge(end_cell)
            return "Cells merged successfully."
        except Exception as e:
            return f"Error merging table cells: {e}"

    def split_table(self, table_index: int, split_row_index: int) -> str:
        """Splits a table into two separate tables at a specified row.

        Args:
            table_index (int): The index of the table.
            split_row_index: The row index where the split should occur.

        Returns:
            A confirmation message indicating the table was split successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            if table_index < 0 or table_index >= len(self.document.tables):
                return "Invalid table index."
            original_table = self.document.tables[table_index]
            if split_row_index < 0 or split_row_index >= len(original_table.rows):
                return "Invalid split row index."

            # Create a new table with the same number of columns
            new_table = self.document.add_table(rows=0, cols=len(original_table.columns))

            # Move rows from the original table to the new table
            for row in original_table.rows[split_row_index:]:
                new_table._tbl.append(row._element)

            # Remove moved rows from the original table
            for row in original_table.rows[split_row_index:]:
                row._element.getparent().remove(row._element)

            return "Table split successfully."
        except Exception as e:
            return f"Error splitting table: {e}"

    def set_page_margins(self, top: float, bottom: float, left: float, right: float) -> str:
        """Sets the page margins for the document.

        Args:
            top (float): The top margin in inches.
            bottom (float): The bottom margin in inches.
            left (float): The left margin in inches.
            right (float): The right margin in inches.

        Returns:
            A confirmation message indicating the margins were set successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            sections = self.document.sections
            for section in sections:
                section.top_margin = Inches(top)
                section.bottom_margin = Inches(bottom)
                section.left_margin = Inches(left)
                section.right_margin = Inches(right)
            return "Page margins set successfully."
        except Exception as e:
            return f"Error setting page margins: {e}"

    def replace_section(self, section_title: str, new_content: str) -> str:
        """Replaces the content of a section based on its title.

        Args:
            section_title (str): The title of the section to replace.
            new_content (str): The new content for the section.

        Returns:
            A confirmation message indicating the section was replaced successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            for paragraph in self.document.paragraphs:
                if paragraph.text.strip() == section_title:
                    next_paragraph = paragraph._element.getnext()
                    while next_paragraph is not None and next_paragraph.tag.endswith('p'):
                        next_paragraph.getparent().remove(next_paragraph)
                        next_paragraph = paragraph._element.getnext()
                    self.document.add_paragraph(new_content)
                    return "Section replaced successfully."
            return "Section title not found."
        except Exception as e:
            return f"Error replacing section: {e}"

    def edit_section_by_keyword(self, keyword: str, new_content: str) -> str:
        """Edits a section of the document based on a keyword.

        Args:
            keyword (str): The keyword to locate the section.
            new_content (str): The new content for the section.

        Returns:
            A confirmation message indicating the section was edited successfully.
        """
        try:
            if not self.document:
                return "No document is currently open."
            for paragraph in self.document.paragraphs:
                if keyword in paragraph.text:
                    paragraph.text = new_content
                    return "Section edited successfully."
            return "Keyword not found."
        except Exception as e:
            return f"Error editing section by keyword: {e}"


manager = WordDocumentManager()


@mcp.tool()
def create_document(file_name: str) -> str:
    """Creates a new Word document.

    Args:
        file_name (str): The name of the file to be created.

    Returns:
        A confirmation message indicating the document was created successfully.
    """
    return manager.create_document(file_name)


@mcp.tool()
def open_document(file_path: str) -> str:
    """Opens an existing Word document.

    Args:
        file_path (str): The path to the document file.

    Returns:
        A confirmation message indicating the document was opened successfully.
    """
    return manager.open_document(file_path)


@mcp.tool()
def save_document(file_path: str) -> str:
    """Saves the currently open Word document.

    Args:
        file_path (str): The path where the document should be saved.

    Returns:
        A confirmation message indicating the document was saved successfully.
    """
    return manager.save_document(file_path)


@mcp.tool()
def save_as_document(new_file_path: str) -> str:
    """Saves the currently open Word document with a new name or in a different location.

    Args:
        new_file_path (str): The new file path for saving the document.

    Returns:
        A confirmation message indicating the document was saved successfully.
    """
    return manager.save_as_document(new_file_path)


@mcp.tool()
def create_document_copy(copy_file_path: str) -> str:
    """Creates a copy of the currently open Word document.

    Args:
        copy_file_path (str): The file path where the copy should be saved.

    Returns:
        A confirmation message indicating the document copy was created successfully.
    """
    return manager.create_document_copy(copy_file_path)


@mcp.tool()
def add_paragraph(text: str) -> str:
    """Adds a paragraph to the document.

    Args:
        text (str): The text content of the paragraph.

    Returns:
        A confirmation message indicating the paragraph was added successfully.
    """
    return manager.add_paragraph(text)


@mcp.tool()
def add_heading(text: str, level: int) -> str:
    """Adds a heading to the document.

    Args:
        text (str): The text content of the heading.
        level (int): The level of the heading (e.g., 1 for main heading, 2 for subheading).

    Returns:
        A confirmation message indicating the heading was added successfully.
    """
    return manager.add_heading(text, level)


@mcp.tool()
def add_table(rows: int, cols: int) -> str:
    """Adds a table to the document.

    Args:
        rows (int): The number of rows in the table.
        cols (int): The number of columns in the table.

    Returns:
        A confirmation message indicating the table was added successfully.
    """
    return manager.add_table(rows, cols)


@mcp.tool()
def add_page_break() -> str:
    """Adds a page break to the document.

    Returns:
        A confirmation message indicating the page break was added successfully.
    """
    return manager.add_page_break()


@mcp.tool()
def get_document_info() -> str:
    """Retrieves metadata about the document, such as the number of pages, paragraphs, and tables.

    Returns:
        A JSON string containing metadata about the document.
    """
    info = manager.get_document_info()
    return json.dumps(info)


@mcp.tool()
def search_text(query: str) -> str:
    """Searches for specific text within the document.

    Args:
        query (str): The text to search for.

    Returns:
        A JSON string of locations where the text was found.
    """
    locations = manager.search_text(query)
    return json.dumps(locations)


@mcp.tool()
def find_and_replace(old_text: str, new_text: str) -> str:
    """Finds specific text in the document and replaces it with new text.

    Args:
        old_text (str): The text to find.
        new_text (str): The text to replace it with.

    Returns:
        A confirmation message indicating the replacement was successful.
    """
    return manager.find_and_replace(old_text, new_text)


@mcp.tool()
def delete_paragraph(paragraph_index: int) -> str:
    """Deletes a specified paragraph from the document.

    Args:
        paragraph_index (int): The index of the paragraph to delete.

    Returns:
        A confirmation message indicating the paragraph was deleted successfully.
    """
    return manager.delete_paragraph(paragraph_index)


@mcp.tool()
def delete_text(text_to_delete: str) -> str:
    """Deletes all instances of specified text from the document.

    Args:
        text_to_delete (str): The text to delete.

    Returns:
        A confirmation message indicating the text was deleted successfully.
    """
    return manager.delete_text(text_to_delete)


@mcp.tool()
def add_table_row(table_index: int) -> str:
    """Adds a new row to an existing table.

    Args:
        table_index (int): The index of the table.

    Returns:
        A confirmation message indicating the row was added successfully.
    """
    return manager.add_table_row(table_index)


@mcp.tool()
def delete_table_row(table_index: int, row_index: int) -> str:
    """Deletes a row from an existing table.

    Args:
        table_index (int): The index of the table.
        row_index (int): The index of the row to delete.

    Returns:
        A confirmation message indicating the row was deleted successfully.
    """
    return manager.delete_table_row(table_index, row_index)


@mcp.tool()
def edit_table_cell(table_index: int, row_index: int, col_index: int, new_content: str) -> str:
    """Edits the content of a specific cell in a table.

    Args:
        table_index (int): The index of the table.
        row_index (int): The index of the row.
        col_index (int): The index of the column.
        new_content (str): The new content for the cell.

    Returns:
        A confirmation message indicating the cell was edited successfully.
    """
    return manager.edit_table_cell(table_index, row_index, col_index, new_content)


@mcp.tool()
def merge_table_cells(table_index: int, start_row: int, start_col: int, end_row: int, end_col: int) -> str:
    """Merges cells in a table.

    Args:
        table_index (int): The index of the table.
        start_row (int): The starting row index for merging.
        start_col (int): The starting column index for merging.
        end_row (int): The ending row index for merging.
        end_col (int): The ending column index for merging.

    Returns:
        A confirmation message indicating the cells were merged successfully.
    """
    return manager.merge_table_cells(table_index, start_row, start_col, end_row, end_col)


@mcp.tool()
def split_table(table_index: int, split_row_index: int) -> str:
    """Splits a table into two separate tables at a specified row.

    Args:
        table_index (int): The index of the table.
        split_row_index (int): The row index where the split should occur.

    Returns:
        A confirmation message indicating the table was split successfully.
    """
    return manager.split_table(table_index, split_row_index)


@mcp.tool()
def set_page_margins(top: float, bottom: float, left: float, right: float) -> str:
    """Sets the page margins for the document.

    Args:
        top (float): The top margin in inches.
        bottom (float): The bottom margin in inches.
        left (float): The left margin in inches.
        right (float): The right margin in inches.

    Returns:
        A confirmation message indicating the margins were set successfully.
    """
    return manager.set_page_margins(top, bottom, left, right)


@mcp.tool()
def replace_section(section_title: str, new_content: str) -> str:
    """Replaces the content of a section based on its title.

    Args:
        section_title (str): The title of the section to replace.
        new_content (str): The new content for the section.

    Returns:
        A confirmation message indicating the section was replaced successfully.
    """
    return manager.replace_section(section_title, new_content)


@mcp.tool()
def edit_section_by_keyword(keyword: str, new_content: str) -> str:
    """Edits a section of the document based on a keyword.

    Args:
        keyword (str): The keyword to locate the section.
        new_content (str): The new content for the section.

    Returns:
        A confirmation message indicating the section was edited successfully.
    """
    return manager.edit_section_by_keyword(keyword, new_content)


if __name__ == "__main__":
    mcp.run()