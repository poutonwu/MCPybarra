import sys
import os
from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import json

# Initialize FastMCP server
mcp = FastMCP("mcp_word_document_automation")

# Global variable to hold the currently opened document
current_document = None

@mcp.tool()
def create_document() -> str:
    """
    Creates a new Word document.

    Returns:
        str: A success message along with the path of the created document.

    Example:
        create_document()
    """
    global current_document
    try:
        current_document = Document()
        file_path = "new_document.docx"
        current_document.save(file_path)
        return json.dumps({"message": "Document created successfully", "file_path": file_path})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def open_document(file_path: str) -> str:
    """
    Opens an existing Word document.

    Args:
        file_path (str): Path to the Word document to open.

    Returns:
        str: A success message.

    Example:
        open_document("example.docx")
    """
    global current_document
    try:
        current_document = Document(file_path)
        return json.dumps({"message": "Document opened successfully"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def save_document() -> str:
    """
    Saves the currently opened Word document.

    Returns:
        str: A success message.

    Example:
        save_document()
    """
    global current_document
    try:
        if current_document:
            file_path = "saved_document.docx"
            current_document.save(file_path)
            return json.dumps({"message": "Document saved successfully", "file_path": file_path})
        else:
            return json.dumps({"error": "No document is currently opened"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def save_as_document(new_file_path: str) -> str:
    """
    Saves the currently opened Word document under a new name.

    Args:
        new_file_path (str): The new file path for saving the document.

    Returns:
        str: A success message.

    Example:
        save_as_document("new_name.docx")
    """
    global current_document
    try:
        if current_document:
            current_document.save(new_file_path)
            return json.dumps({"message": "Document saved successfully", "file_path": new_file_path})
        else:
            return json.dumps({"error": "No document is currently opened"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def create_document_copy(copy_file_path: str) -> str:
    """
    Creates a copy of the currently opened Word document.

    Args:
        copy_file_path (str): The file path for the copied document.

    Returns:
        str: A success message.

    Example:
        create_document_copy("copy_name.docx")
    """
    global current_document
    try:
        if current_document:
            current_document.save(copy_file_path)
            return json.dumps({"message": "Document copied successfully", "file_path": copy_file_path})
        else:
            return json.dumps({"error": "No document is currently opened"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def add_paragraph(text: str) -> str:
    """
    Adds a paragraph to the document.

    Args:
        text (str): The text content of the paragraph.

    Returns:
        str: A success message.

    Example:
        add_paragraph("This is a paragraph.")
    """
    global current_document
    try:
        if current_document:
            current_document.add_paragraph(text)
            return json.dumps({"message": "Paragraph added successfully"})
        else:
            return json.dumps({"error": "No document is currently opened"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def add_heading(text: str, level: int) -> str:
    """
    Adds a heading to the document.

    Args:
        text (str): The text content of the heading.
        level (int): Heading level (e.g., 1 for H1, 2 for H2).

    Returns:
        str: A success message.

    Example:
        add_heading("Chapter 1", 1)
    """
    global current_document
    try:
        if current_document:
            current_document.add_heading(text, level=level)
            return json.dumps({"message": "Heading added successfully"})
        else:
            return json.dumps({"error": "No document is currently opened"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def add_table(rows: int, columns: int) -> str:
    """
    Adds a table to the document.

    Args:
        rows (int): Number of rows in the table.
        columns (int): Number of columns in the table.

    Returns:
        str: A success message.

    Example:
        add_table(3, 4)
    """
    global current_document
    try:
        if current_document:
            current_document.add_table(rows=rows, cols=columns)
            return json.dumps({"message": "Table added successfully"})
        else:
            return json.dumps({"error": "No document is currently opened"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def add_page_break() -> str:
    """
    Adds a page break to the document.

    Returns:
        str: A success message.

    Example:
        add_page_break()
    """
    global current_document
    try:
        if current_document:
            current_document.add_page_break()
            return json.dumps({"message": "Page break added successfully"})
        else:
            return json.dumps({"error": "No document is currently opened"})
    except Exception as e:
        return json.dumps({"error": str(e)})

if __name__ == "__main__":
    # Ensure UTF-8 encoding for standard output
    sys.stdout.reconfigure(encoding='utf-8')
    mcp.run()