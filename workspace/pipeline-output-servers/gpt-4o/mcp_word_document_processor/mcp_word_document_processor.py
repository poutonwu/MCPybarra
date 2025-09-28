import sys
import os
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from mcp.server.fastmcp import FastMCP

# Initialize MCP Server
mcp = FastMCP("mcp_word_document_processor")

@mcp.tool()
def create_document(title: str, author: str, subject: str) -> str:
    """
    Creates a new Word document and sets metadata such as title, author, and subject.

    Args:
        title (str): The title of the document.
        author (str): The author of the document.
        subject (str): The subject of the document.

    Returns:
        str: A success message with the document's file path.

    Example:
        create_document(title="My Document", author="John Doe", subject="Sample Subject")
    """
    try:
        document = Document()
        core_properties = document.core_properties
        core_properties.title = title
        core_properties.author = author
        core_properties.subject = subject

        file_path = f"{title.replace(' ', '_')}.docx"
        document.save(file_path)

        return json.dumps({"message": "Document created successfully", "file_path": file_path})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def get_document_text(file_path: str) -> str:
    """
    Extracts all text content from an existing Word document.

    Args:
        file_path (str): The path to the document file.

    Returns:
        str: A string containing the full text of the document.

    Example:
        get_document_text(file_path="example.docx")
    """
    try:
        document = Document(file_path)
        text = "\n".join([p.text for p in document.paragraphs])
        return json.dumps({"text": text})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def add_paragraph(file_path: str, text: str) -> str:
    """
    Adds a paragraph of text to the document.

    Args:
        file_path (str): The path to the document file.
        text (str): The paragraph text to add.

    Returns:
        str: A success message confirming the addition.

    Example:
        add_paragraph(file_path="example.docx", text="This is a new paragraph.")
    """
    try:
        document = Document(file_path)
        document.add_paragraph(text)
        document.save(file_path)
        return json.dumps({"message": "Paragraph added successfully"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def add_heading(file_path: str, text: str, level: int) -> str:
    """
    Adds a heading of a specified level to the document.

    Args:
        file_path (str): The path to the document file.
        text (str): The heading text.
        level (int): The heading level (1-6).

    Returns:
        str: A success message confirming the addition.

    Example:
        add_heading(file_path="example.docx", text="Chapter 1", level=1)
    """
    try:
        document = Document(file_path)
        document.add_heading(text, level=level)
        document.save(file_path)
        return json.dumps({"message": "Heading added successfully"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def create_custom_style(file_path: str, style_name: str, font_name: str, font_size: int, bold: bool, italic: bool) -> str:
    """
    Creates a custom text style for the document.

    Args:
        file_path (str): The path to the document file.
        style_name (str): The name of the custom style.
        font_name (str): The font name for the style.
        font_size (int): The font size.
        bold (bool): Whether the text is bold.
        italic (bool): Whether the text is italicized.

    Returns:
        str: A success message confirming the style creation.

    Example:
        create_custom_style(file_path="example.docx", style_name="CustomStyle", font_name="Arial", font_size=12, bold=True, italic=False)
    """
    try:
        document = Document(file_path)
        styles = document.styles
        new_style = styles.add_style(style_name, WD_PARAGRAPH_ALIGNMENT.LEFT)
        font = new_style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold
        font.italic = italic
        document.save(file_path)
        return json.dumps({"message": "Custom style created successfully"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def format_text(file_path: str, start: int, end: int, style_name: str) -> str:
    """
    Formats a specific range of text in the document with a custom style.

    Args:
        file_path (str): The path to the document file.
        start (int): The start index of the text range.
        end (int): The end index of the text range.
        style_name (str): The name of the custom style to apply.

    Returns:
        str: A success message confirming the formatting.

    Example:
        format_text(file_path="example.docx", start=0, end=10, style_name="CustomStyle")
    """
    try:
        document = Document(file_path)
        paragraphs = document.paragraphs[start:end]
        for paragraph in paragraphs:
            paragraph.style = style_name
        document.save(file_path)
        return json.dumps({"message": "Text formatted successfully"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def protect_document(file_path: str, password: str) -> str:
    """
    Adds password protection to the document.

    Args:
        file_path (str): The path to the document file.
        password (str): The password to protect the document.

    Returns:
        str: A success message confirming the protection.

    Example:
        protect_document(file_path="example.docx", password="securepassword")
    """
    try:
        # Note: python-docx does not support password protection directly.
        # Implementing this feature would require external libraries or tools.
        return json.dumps({"error": "Password protection not implemented"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def add_footnote_to_document(file_path: str, text: str) -> str:
    """
    Adds a footnote to the document.

    Args:
        file_path (str): The path to the document file.
        text (str): The footnote text.

    Returns:
        str: A success message confirming the addition.

    Example:
        add_footnote_to_document(file_path="example.docx", text="This is a footnote.")
    """
    try:
        # Note: python-docx does not support footnotes directly.
        # Implementing this feature would require external libraries or tools.
        return json.dumps({"error": "Footnote addition not implemented"})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def get_paragraph_text_from_document(file_path: str, paragraph_index: int) -> str:
    """
    Retrieves the text of a specific paragraph by index.

    Args:
        file_path (str): The path to the document file.
        paragraph_index (int): The index of the paragraph.

    Returns:
        str: A string containing the text of the specified paragraph.

    Example:
        get_paragraph_text_from_document(file_path="example.docx", paragraph_index=0)
    """
    try:
        document = Document(file_path)
        text = document.paragraphs[paragraph_index].text
        return json.dumps({"text": text})
    except Exception as e:
        return json.dumps({"error": str(e)})

@mcp.tool()
def find_text_in_document(file_path: str, search_text: str) -> str:
    """
    Searches for a specific text in the document and returns its occurrences.

    Args:
        file_path (str): The path to the document file.
        search_text (str): The text to search for.

    Returns:
        str: A list of indices where the text occurs.

    Example:
        find_text_in_document(file_path="example.docx", search_text="sample")
    """
    try:
        document = Document(file_path)
        occurrences = []
        for i, paragraph in enumerate(document.paragraphs):
            if search_text in paragraph.text:
                occurrences.append(i)
        return json.dumps({"occurrences": occurrences})
    except Exception as e:
        return json.dumps({"error": str(e)})

if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    mcp.run("stdio")