import sys
from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.shared import Inches
import os

# Initialize FastMCP server
mcp = FastMCP("mcp_word_document_processor")

# Global variable to hold the current document
current_doc = None

@mcp.tool()
def create_document() -> str:
    """
    Creates a new Word document.

    Returns:
        A string indicating the document ID or file path of the newly created document.
    """
    global current_doc
    current_doc = Document()
    return "New document created successfully."

@mcp.tool()
def open_document(document_path: str) -> str:
    """
    Opens an existing Word document.

    Args:
        document_path: Path to the document file.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    try:
        if not os.path.exists(document_path):
            return "Document path does not exist."
        current_doc = Document(document_path)
        return f"Document opened successfully from {document_path}."
    except Exception as e:
        return f"Failed to open document: {str(e)}"

@mcp.tool()
def save_document() -> str:
    """
    Saves the currently open document.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            current_doc.save("current_document.docx")
            return "Document saved successfully as current_document.docx."
        except Exception as e:
            return f"Failed to save document: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def save_as_document(new_path: str) -> str:
    """
    Saves the currently open document with a new name or path.

    Args:
        new_path: New file path for the document.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            current_doc.save(new_path)
            return f"Document saved successfully as {new_path}."
        except Exception as e:
            return f"Failed to save document: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def create_document_copy(copy_path: str) -> str:
    """
    Creates a copy of the currently open document.

    Args:
        copy_path: Path for the new copy.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            current_doc.save(copy_path)
            return f"Document copy created successfully at {copy_path}."
        except Exception as e:
            return f"Failed to create document copy: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def add_paragraph(text: str) -> str:
    """
    Adds a paragraph to the document.

    Args:
        text: The text content of the paragraph.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            current_doc.add_paragraph(text)
            return "Paragraph added successfully."
        except Exception as e:
            return f"Failed to add paragraph: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def add_heading(text: str, level: int = 1) -> str:
    """
    Adds a heading to the document.

    Args:
        text: The text content of the heading.
        level: Heading level (e.g., 1 for top-level).

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            if level < 1 or level > 9:
                return "Invalid heading level (must be between 1 and 9)."
            current_doc.add_heading(text, level)
            return "Heading added successfully."
        except Exception as e:
            return f"Failed to add heading: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def add_table(rows: int, columns: int) -> str:
    """
    Adds a table to the document.

    Args:
        rows: Number of rows.
        columns: Number of columns.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            if rows <= 0 or columns <= 0:
                return "Rows and columns must be positive integers."
            current_doc.add_table(rows, columns)
            return "Table added successfully."
        except Exception as e:
            return f"Failed to add table: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def add_page_break() -> str:
    """
    Adds a page break to the document.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            current_doc.add_page_break()
            return "Page break added successfully."
        except Exception as e:
            return f"Failed to add page break: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def get_document_info() -> str:
    """
    Retrieves metadata about the document (e.g., word count, author).

    Returns:
        A JSON object containing document metadata.
    """
    global current_doc
    if current_doc:
        try:
            info = {
                "author": current_doc.core_properties.author,
                "created": str(current_doc.core_properties.created),
                "modified": str(current_doc.core_properties.modified),
                "word_count": sum(len(p.text.split()) for p in current_doc.paragraphs)
            }
            return str(info)
        except Exception as e:
            return f"Failed to retrieve document info: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def search_text(query: str) -> str:
    """
    Searches for text in the document.

    Args:
        query: The text to search for.

    Returns:
        A list of matches with their positions.
    """
    global current_doc
    if current_doc:
        try:
            matches = []
            for i, paragraph in enumerate(current_doc.paragraphs):
                if query in paragraph.text:
                    matches.append({"paragraph_index": i, "text": paragraph.text})
            return str(matches)
        except Exception as e:
            return f"Failed to search text: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def find_and_replace(find_text: str, replace_text: str) -> str:
    """
    Finds and replaces text in the document.

    Args:
        find_text: The text to find.
        replace_text: The replacement text.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            for paragraph in current_doc.paragraphs:
                if find_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(find_text, replace_text)
            return "Text replaced successfully."
        except Exception as e:
            return f"Failed to replace text: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def delete_paragraph(paragraph_index: int) -> str:
    """
    Deletes a paragraph from the document.

    Args:
        paragraph_index: Index of the paragraph to delete.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            if 0 <= paragraph_index < len(current_doc.paragraphs):
                p = current_doc.paragraphs[paragraph_index]
                p.clear()
                return "Paragraph deleted successfully."
            else:
                return "Invalid paragraph index."
        except Exception as e:
            return f"Failed to delete paragraph: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def delete_text(text: str) -> str:
    """
    Deletes specified text from the document.

    Args:
        text: The text to delete.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            for paragraph in current_doc.paragraphs:
                if text in paragraph.text:
                    paragraph.text = paragraph.text.replace(text, "")
            return "Text deleted successfully."
        except Exception as e:
            return f"Failed to delete text: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def add_table_row(table_index: int) -> str:
    """
    Adds a row to an existing table.

    Args:
        table_index: Index of the table.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            if 0 <= table_index < len(current_doc.tables):
                table = current_doc.tables[table_index]
                table.add_row()
                return "Row added successfully."
            else:
                return "Invalid table index."
        except Exception as e:
            return f"Failed to add row: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def delete_table_row(table_index: int, row_index: int) -> str:
    """
    Deletes a row from an existing table.

    Args:
        table_index: Index of the table.
        row_index: Index of the row to delete.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            if 0 <= table_index < len(current_doc.tables):
                table = current_doc.tables[table_index]
                if 0 <= row_index < len(table.rows):
                    table._tbl.remove(table.rows[row_index]._tr)
                    return "Row deleted successfully."
                else:
                    return "Invalid row index."
            else:
                return "Invalid table index."
        except Exception as e:
            return f"Failed to delete row: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def edit_table_cell(table_index: int, row_index: int, column_index: int, new_text: str) -> str:
    """
    Edits the content of a table cell.

    Args:
        table_index: Index of the table.
        row_index: Row index of the cell.
        column_index: Column index of the cell.
        new_text: New content for the cell.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            if 0 <= table_index < len(current_doc.tables):
                table = current_doc.tables[table_index]
                if 0 <= row_index < len(table.rows) and 0 <= column_index < len(table.columns):
                    cell = table.cell(row_index, column_index)
                    cell.text = new_text
                    return "Cell edited successfully."
                else:
                    return "Invalid row or column index."
            else:
                return "Invalid table index."
        except Exception as e:
            return f"Failed to edit cell: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def merge_table_cells(table_index: int, start_row: int, start_column: int, end_row: int, end_column: int) -> str:
    """
    Merges specified cells in a table.

    Args:
        table_index: Index of the table.
        start_row: Starting row index.
        start_column: Starting column index.
        end_row: Ending row index.
        end_column: Ending column index.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            if 0 <= table_index < len(current_doc.tables):
                table = current_doc.tables[table_index]
                if 0 <= start_row < len(table.rows) and 0 <= start_column < len(table.columns) and \
                   0 <= end_row < len(table.rows) and 0 <= end_column < len(table.columns):
                    cell = table.cell(start_row, start_column)
                    cell.merge(table.cell(end_row, end_column))
                    return "Cells merged successfully."
                else:
                    return "Invalid row or column indices."
            else:
                return "Invalid table index."
        except Exception as e:
            return f"Failed to merge cells: {str(e)}"
    else:
        return "No document is currently open."

@mcp.tool()
def set_page_margins(left: float, right: float, top: float, bottom: float) -> str:
    """
    Sets the page margins for the document.

    Args:
        left: Left margin in inches.
        right: Right margin in inches.
        top: Top margin in inches.
        bottom: Bottom margin in inches.

    Returns:
        A string indicating success or failure.
    """
    global current_doc
    if current_doc:
        try:
            sections = current_doc.sections
            for section in sections:
                section.left_margin = Inches(left)
                section.right_margin = Inches(right)
                section.top_margin = Inches(top)
                section.bottom_margin = Inches(bottom)
            return "Page margins set successfully."
        except Exception as e:
            return f"Failed to set page margins: {str(e)}"
    else:
        return "No document is currently open."

if __name__ == "__main__":
    sys.stdout.reconfigure(encoding='utf-8')
    mcp.run()