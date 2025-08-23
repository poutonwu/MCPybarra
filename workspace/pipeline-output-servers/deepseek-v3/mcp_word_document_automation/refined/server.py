from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import sys
import os
from mcp.server.fastmcp import FastMCP

# Initialize FastMCP server
mcp = FastMCP("mcp_word_document_automation")

@mcp.tool()
def create_document(title: str, author: str = "Unknown", keywords: list = None) -> str:
    """
    Creates a new Word document and sets metadata such as title, author, and keywords.

    Args:
        title: The title of the document.
        author: The author of the document. Defaults to "Unknown".
        keywords: A list of keywords for the document. Defaults to an empty list.

    Returns:
        A string containing the path or identifier of the newly created document.
    """
    try:
        os.makedirs("workspace", exist_ok=True)
        document = Document()
        document.core_properties.title = title
        document.core_properties.author = author
        if keywords:
            document.core_properties.keywords = ", ".join(keywords)
        file_path = f"workspace/{title}.docx"
        document.save(file_path)
        return file_path
    except Exception as e:
        raise Exception(f"Failed to create document: {str(e)}")

@mcp.tool()
def get_document_text(document_id: str) -> str:
    """
    Extracts the full text content of a specified Word document.

    Args:
        document_id: The identifier or path of the document.

    Returns:
        A string containing the full text of the document.
    """
    try:
        document = Document(document_id)
        full_text = "\n".join([para.text for para in document.paragraphs])
        return full_text
    except Exception as e:
        raise Exception(f"Failed to extract text from document: {str(e)}")

@mcp.tool()
def add_paragraph(document_id: str, text: str) -> bool:
    """
    Adds a paragraph of text to a specified Word document.

    Args:
        document_id: The identifier or path of the document.
        text: The text to add as a paragraph.

    Returns:
        A boolean indicating success or failure.
    """
    try:
        document = Document(document_id)
        document.add_paragraph(text)
        document.save(document_id)
        return True
    except Exception as e:
        raise Exception(f"Failed to add paragraph: {str(e)}")

@mcp.tool()
def add_heading(document_id: str, text: str, level: int) -> bool:
    """
    Adds a heading to a specified Word document at a specified level (e.g., H1, H2).

    Args:
        document_id: The identifier or path of the document.
        text: The text of the heading.
        level: The heading level (1 for H1, 2 for H2, etc.).

    Returns:
        A boolean indicating success or failure.
    """
    try:
        if level < 1 or level > 9:
            raise ValueError("Heading level must be between 1 and 9")
        document = Document(document_id)
        document.add_heading(text, level)
        document.save(document_id)
        return True
    except Exception as e:
        raise Exception(f"Failed to add heading: {str(e)}")

@mcp.tool()
def create_custom_style(style_name: str, font: str, size: int, color: str) -> bool:
    """
    Creates a custom text style (e.g., font, size, color) for use in the document.

    Args:
        style_name: The name of the custom style.
        font: The font name.
        size: The font size.
        color: The font color in HEX format (e.g., "#FF0000").

    Returns:
        A boolean indicating success or failure.
    """
    try:
        document = Document()
        styles = document.styles
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = font
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        document.save("workspace/styles.docx")
        return True
    except Exception as e:
        raise Exception(f"Failed to create custom style: {str(e)}")

@mcp.tool()
def format_text(document_id: str, start_pos: int, end_pos: int, format_options: dict) -> bool:
    """
    Applies formatting (e.g., bold, italic) to a specified range of text in the document.

    Args:
        document_id: The identifier or path of the document.
        start_pos: The starting position of the text range.
        end_pos: The ending position of the text range.
        format_options: A dictionary of formatting options (e.g., {"bold": True, "italic": False}).

    Returns:
        A boolean indicating success or failure.
    """
    try:
        if start_pos > end_pos:
            raise ValueError("Start position cannot be greater than end position")
        
        document = Document(document_id)
        for para in document.paragraphs:
            if start_pos <= len(para.text) <= end_pos:
                for run in para.runs:
                    if "bold" in format_options:
                        run.bold = format_options["bold"]
                    if "italic" in format_options:
                        run.italic = format_options["italic"]
                    if "underline" in format_options:
                        run.underline = format_options["underline"]
        document.save(document_id)
        return True
    except Exception as e:
        raise Exception(f"Failed to format text: {str(e)}")

@mcp.tool()
def protect_document(document_id: str, password: str) -> bool:
    """
    Sets password protection for a specified Word document.

    Args:
        document_id: The identifier or path of the document.
        password: The password to set for the document.

    Returns:
        A boolean indicating success or failure.
    """
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(document_id)
        doc.SaveAs(document_id, Password=password)
        doc.Close()
        word.Quit()
        return True
    except ImportError:
        raise Exception("pywin32 package is required for document protection")
    except Exception as e:
        raise Exception(f"Failed to protect document: {str(e)}")

@mcp.tool()
def add_footnote_to_document(document_id: str, text: str) -> bool:
    """
    Adds a footnote to a specified Word document.

    Args:
        document_id: The identifier or path of the document.
        text: The text of the footnote.

    Returns:
        A boolean indicating success or failure.
    """
    try:
        document = Document(document_id)
        document.add_footnote(text)
        document.save(document_id)
        return True
    except Exception as e:
        raise Exception(f"Failed to add footnote: {str(e)}")

@mcp.tool()
def get_paragraph_text_from_document(document_id: str, paragraph_index: int) -> str:
    """
    Extracts the text of a specific paragraph from a Word document.

    Args:
        document_id: The identifier or path of the document.
        paragraph_index: The index of the paragraph to extract.

    Returns:
        A string containing the text of the specified paragraph.
    """
    try:
        document = Document(document_id)
        if 0 <= paragraph_index < len(document.paragraphs):
            return document.paragraphs[paragraph_index].text
        else:
            raise IndexError("Paragraph index out of range")
    except Exception as e:
        raise Exception(f"Failed to get paragraph text: {str(e)}")

@mcp.tool()
def find_text_in_document(document_id: str, search_text: str) -> list:
    """
    Searches for a specified text string in a Word document and returns its positions.

    Args:
        document_id: The identifier or path of the document.
        search_text: The text to search for.

    Returns:
        A list of dictionaries, each containing the start and end positions of the found text.
    """
    try:
        document = Document(document_id)
        positions = []
        for para in document.paragraphs:
            if search_text in para.text:
                start_pos = para.text.find(search_text)
                end_pos = start_pos + len(search_text)
                positions.append({"start": start_pos, "end": end_pos})
        return positions
    except Exception as e:
        raise Exception(f"Failed to find text in document: {str(e)}")

@mcp.tool()
def add_table(document_id: str, rows: int, cols: int, data: list) -> bool:
    """
    Adds a table to a specified Word document.

    Args:
        document_id: The identifier or path of the document.
        rows: The number of rows in the table.
        cols: The number of columns in the table.
        data: The data to populate the table.

    Returns:
        A boolean indicating success or failure.
    """
    try:
        if len(data) != rows or any(len(row) != cols for row in data):
            raise ValueError("Data dimensions must match rows and columns")
            
        document = Document(document_id)
        table = document.add_table(rows=rows, cols=cols)
        for i, row in enumerate(data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = str(cell)
        document.save(document_id)
        return True
    except Exception as e:
        raise Exception(f"Failed to add table: {str(e)}")

@mcp.tool()
def add_image(document_id: str, image_path: str) -> bool:
    """
    Adds an image to a specified Word document.

    Args:
        document_id: The identifier or path of the document.
        image_path: The path to the image file.

    Returns:
        A boolean indicating success or failure.
    """
    try:
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Image file not found: {image_path}")
            
        document = Document(document_id)
        document.add_picture(image_path)
        document.save(document_id)
        return True
    except Exception as e:
        raise Exception(f"Failed to add image: {str(e)}")

@mcp.tool()
def add_page_break(document_id: str) -> bool:
    """
    Adds a page break to a specified Word document.

    Args:
        document_id: The identifier or path of the document.

    Returns:
        A boolean indicating success or failure.
    """
    try:
        document = Document(document_id)
        document.add_page_break()
        document.save(document_id)
        return True
    except Exception as e:
        raise Exception(f"Failed to add page break: {str(e)}")

@mcp.tool()
def add_header_footer(document_id: str, text: str, is_header: bool) -> bool:
    """
    Adds a header or footer to a specified Word document.

    Args:
        document_id: The identifier or path of the document.
        text: The text to add to the header or footer.
        is_header: True for header, False for footer.

    Returns:
        A boolean indicating success or failure.
    """
    try:
        document = Document(document_id)
        if is_header:
            section = document.sections[0]
            header = section.header
            header.paragraphs[0].text = text
        else:
            section = document.sections[0]
            footer = section.footer
            footer.paragraphs[0].text = text
        document.save(document_id)
        return True
    except Exception as e:
        raise Exception(f"Failed to add header/footer: {str(e)}")

@mcp.tool()
def convert_to_pdf(document_id: str) -> str:
    """
    Converts a specified Word document to PDF format.

    Args:
        document_id: The identifier or path of the document.

    Returns:
        A string containing the path or identifier of the converted PDF.
    """
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(document_id)
        pdf_path = document_id.replace(".docx", ".pdf")
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()
        return pdf_path
    except ImportError:
        raise Exception("pywin32 package is required for PDF conversion")
    except Exception as e:
        raise Exception(f"Failed to convert to PDF: {str(e)}")

if __name__ == "__main__":
    sys.stdout.reconfigure(encoding='utf-8')
    mcp.run()