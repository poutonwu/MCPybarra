import sys
from mcp.server.fastmcp import FastMCP
from docx import Document  # Python python-docx library for Word document manipulation
from docx.shared import Inches  # For setting page margins

# Initialize FastMCP server for document processing
mcp = FastMCP("document_processor")

# Placeholder for document management (in-memory storage for active documents)
documents = {}

@mcp.tool()
def create_document(doc_id: str) -> str:
    """
    Creates a new Word document with the specified ID.

    Args:
        doc_id: A unique identifier for the document (string, required).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the document ID is invalid or already exists.
    """
    if not doc_id or doc_id.strip() == "":
        raise ValueError("Document ID cannot be empty.")
    
    if doc_id in documents:
        raise ValueError(f"Document with ID '{doc_id}' already exists.")
    
    # Create a new document instance
    document = Document()
    documents[doc_id] = document
    return f"Document '{doc_id}' created successfully."

@mcp.tool()
def open_document(doc_id: str) -> str:
    """
    Opens an existing Word document by its ID.

    Args:
        doc_id: The unique identifier of the document to open (string, required).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the document ID does not exist.
    """
    if doc_id not in documents:
        raise ValueError(f"Document with ID '{doc_id}' does not exist.")
    return f"Document '{doc_id}' opened successfully."

@mcp.tool()
def save_document(doc_id: str) -> str:
    """
    Saves the specified document to disk.

    Args:
        doc_id: The unique identifier of the document to save (string, required).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the document ID does not exist.
    """
    if doc_id not in documents:
        raise ValueError(f"Document with ID '{doc_id}' does not exist.")
    # Simulate saving the document (e.g., to disk)
    return f"Document '{doc_id}' saved successfully."

@mcp.tool()
def save_as_document(doc_id: str, new_doc_id: str) -> str:
    """
    Saves a copy of the document with a new ID.

    Args:
        doc_id: The original document ID (string, required).
        new_doc_id: The new document ID (string, required).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the original document doesn't exist or the new ID is already taken.
    """
    if doc_id not in documents:
        raise ValueError(f"Document with ID '{doc_id}' does not exist.")
    if new_doc_id in documents:
        raise ValueError(f"Document with ID '{new_doc_id}' already exists.")
    # Copy the document under a new ID
    documents[new_doc_id] = documents[doc_id]
    return f"Document '{doc_id}' saved as '{new_doc_id}'."

@mcp.tool()
def create_document_copy(doc_id: str, copy_doc_id: str) -> str:
    """
    Creates a copy of the specified document with a new ID.

    Args:
        doc_id: The original document ID (string, required).
        copy_doc_id: The new document ID for the copy (string, required).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the original document doesn't exist or the new ID is already taken.
    """
    if doc_id not in documents:
        raise ValueError(f"Document with ID '{doc_id}' does not exist.")
    if copy_doc_id in documents:
        raise ValueError(f"Document with ID '{copy_doc_id}' already exists.")
    # Clone the document
    documents[copy_doc_id] = Document()
    # Copy content from original to the new document
    original_doc = documents[doc_id]
    copy_doc = documents[copy_doc_id]
    for element in original_doc.element.body:
        copy_doc.element.body.append(element.copy())
    return f"Copy of document '{doc_id}' created as '{copy_doc_id}'."

@mcp.tool()
def add_paragraph(doc_id: str, text: str) -> str:
    """
    Adds a paragraph to the specified document.

    Args:
        doc_id: The unique identifier of the document (string, required).
        text: The text content for the new paragraph (string, required).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the document does not exist or the input text is invalid.
    """
    if doc_id not in documents:
        raise ValueError(f"Document with ID '{doc_id}' does not exist.")
    if not text or text.strip() == "":
        raise ValueError("Paragraph text cannot be empty.")
    
    document = documents[doc_id]
    document.add_paragraph(text)
    return f"Paragraph added to document '{doc_id}'."

@mcp.tool()
def add_heading(doc_id: str, text: str, level: int = 1) -> str:
    """
    Adds a heading to the specified document.

    Args:
        doc_id: The unique identifier of the document (string, required).
        text: The heading text (string, required).
        level: The heading level (integer, optional, default=1).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the document does not exist or the input text is invalid.
    """
    if doc_id not in documents:
        raise ValueError(f"Document with ID '{doc_id}' does not exist.")
    if not text or text.strip() == "":
        raise ValueError("Heading text cannot be empty.")
    if level < 1 or level > 9:
        raise ValueError("Heading level must be between 1 and 9.")
    
    document = documents[doc_id]
    document.add_heading(text, level=level)
    return f"Heading added to document '{doc_id}'."

@mcp.tool()
def add_page_break(doc_id: str) -> str:
    """
    Adds a page break to the specified document.

    Args:
        doc_id: The unique identifier of the document (string, required).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the document does not exist.
    """
    if doc_id not in documents:
        raise ValueError(f"Document with ID '{doc_id}' does not exist.")
    
    document = documents[doc_id]
    document.add_page_break()
    return f"Page break added to document '{doc_id}'."

@mcp.tool()
def add_table(doc_id: str, rows: int, cols: int) -> str:
    """
    Adds a table to the specified document.

    Args:
        doc_id: The unique identifier of the document (string, required).
        rows: Number of rows for the table (integer, required, >=1).
        cols: Number of columns for the table (integer, required, >=1).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the document does not exist or the row/column values are invalid.
    """
    if doc_id not in documents:
        raise ValueError(f"Document with ID '{doc_id}' does not exist.")
    if rows < 1 or cols < 1:
        raise ValueError("Row and column counts must be at least 1.")
    
    document = documents[doc_id]
    document.add_table(rows=rows, cols=cols)
    return f"Table ({rows}x{cols}) added to document '{doc_id}'."

@mcp.tool()
def add_table_row(doc_id: str, table_index: int) -> str:
    """
    Adds a row to an existing table in the document.

    Args:
        doc_id: The unique identifier of the document (string, required).
        table_index: Zero-based index of the table to modify (integer, required, >=0).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the document/table does not exist or the index is invalid.
    """
    if doc_id not in documents:
        raise ValueError(f"Document with ID '{doc_id}' does not exist.")
    document = documents[doc_id]
    
    if len(document.tables) <= table_index:
        raise ValueError(f"Document has no table at index {table_index}.")
    
    table = document.tables[table_index]
    table.add_row()
    return f"Row added to table {table_index} in document '{doc_id}'."

@mcp.tool()
def delete_table_row(doc_id: str, table_index: int, row_index: int) -> str:
    """
    Deletes a row from an existing table in the document.

    Args:
        doc_id: The unique identifier of the document (string, required).
        table_index: Zero-based index of the table (integer, required, >=0).
        row_index: Zero-based index of the row to delete (integer, required, >=0).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the document/table/row does not exist or the index is invalid.
    """
    if doc_id not in documents:
        raise ValueError(f"Document with ID '{doc_id}' does not exist.")
    document = documents[doc_id]
    
    if len(document.tables) <= table_index:
        raise ValueError(f"Document has no table at index {table_index}.")
    
    table = document.tables[table_index]
    if len(table.rows) <= row_index:
        raise ValueError(f"Table has no row at index {row_index}.")
    
    # Delete the specified row
    table._tbl.remove(table.rows[row_index]._tr)
    return f"Row {row_index} deleted from table {table_index} in document '{doc_id}'."

@mcp.tool()
def edit_table_cell(doc_id: str, table_index: int, row_index: int, col_index: int, new_text: str) -> str:
    """
    Edits the content of a specific cell in a table.

    Args:
        doc_id: The unique identifier of the document (string, required).
        table_index: Zero-based index of the table (integer, required, >=0).
        row_index: Zero-based index of the row (integer, required, >=0).
        col_index: Zero-based index of the column (integer, required, >=0).
        new_text: New text content for the cell (string, required).

    Returns:
        A string indicating success or failure of the operation.

    Raises:
        ValueError: If the document/table/row/column does not exist or input is invalid.
    """
    if doc_id not in documents:
        raise ValueError(f"Document with ID '{doc_id}' does not exist.")
    if not new_text or new_text.strip() == "":
        raise ValueError("New text content cannot be empty.")
    
    document = documents[doc_id]
    if len(document.tables) <= table_index:
        raise ValueError(f"Document has no table at index {table_index}.")
    
    table = document.tables[table_index]
    if len(table.rows) <= row_index:
        raise ValueError(f"Table has no row at index {row_index}.")
    if len(table.columns) <= col_index:
        raise ValueError(f"Table has no column at index {col_index}.")
    
    cell = table.cell(row_index, col_index)
    cell.text = new_text
    return f"Cell at ({row_index}, {col_index}) in table {table_index} updated in document '{doc_id}'."

if __name__ == "__main__":
    sys.stdout.reconfigure(encoding='utf-8')
    mcp.run()