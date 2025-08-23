import sys
from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
import io
import os
try:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    pass

# 初始化 FastMCP 服务器
mcp = FastMCP("doc_processor")

# 文档存储路径
DOC_STORAGE_PATH = "./documents"
os.makedirs(DOC_STORAGE_PATH, exist_ok=True)

@mcp.tool()
def create_document(document_id: str, title: str = "", author: str = "", subject: str = "") -> str:
    """
    创建一个新的Word文档并设置元数据。
    
    Args:
        document_id: 文档的唯一标识符（必填）。
        title: 文档标题（可选）。
        author: 文档作者（可选）。
        subject: 文档主题（可选）。
    
    Returns:
        一个字符串，表示文档创建成功或失败的消息。
    
    示例:
        create_document(document_id="doc1", title="报告", author="张三", subject="年度报告")
    """
    # 创建新文档
    doc = Document()
    
    # 设置文档元数据
    doc.core_properties.title = title
    doc.core_properties.author = author
    doc.core_properties.subject = subject
    
    # 保存文档
    doc.save(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    return f"文档 {document_id} 已成功创建。"

@mcp.tool()
def get_document_text(document_id: str) -> str:
    """
    提取文档的全文内容。
    
    Args:
        document_id: 文档的唯一标识符（必填）。
    
    Returns:
        一个字符串，包含文档的全文内容。
    
    示例:
        get_document_text(document_id="doc1")
    """
    # 加载文档
    doc = Document(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    # 提取所有段落文本
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    return text

@mcp.tool()
def add_paragraph(document_id: str, text: str, style: str = "") -> str:
    """
    向文档添加段落文本。
    
    Args:
        document_id: 文档的唯一标识符（必填）。
        text: 要添加的段落文本（必填）。
        style: 段落样式名称（可选）。
    
    Returns:
        一个字符串，表示操作成功或失败的消息。
    
    示例:
        add_paragraph(document_id="doc1", text="这是一个新段落。", style="Normal")
    """
    # 加载文档
    doc = Document(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    # 添加段落
    paragraph = doc.add_paragraph(text=text)
    
    # 应用样式
    if style:
        try:
            paragraph.style = doc.styles[style]
        except KeyError:
            raise ValueError(f"样式 '{style}' 不存在。")
    
    # 保存文档
    doc.save(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    return f"段落已成功添加到文档 {document_id}。"

@mcp.tool()
def add_heading(document_id: str, text: str, level: int = 1) -> str:
    """
    向文档添加各级标题。
    
    Args:
        document_id: 文档的唯一标识符（必填）。
        text: 要添加的标题文本（必填）。
        level: 标题级别（1-9），1为最高级标题（必填）。
    
    Returns:
        一个字符串，表示操作成功或失败的消息。
    
    示例:
        add_heading(document_id="doc1", text="第一章 引言", level=1)
    """
    # 参数验证
    if not 1 <= level <= 9:
        raise ValueError(f"标题级别必须在1-9之间，当前值: {level}")
    
    # 加载文档
    doc = Document(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    # 添加标题
    heading = doc.add_heading(text=text, level=level)
    
    # 保存文档
    doc.save(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    return f"标题 '{text}' 已成功添加到文档 {document_id}。"

@mcp.tool()
def create_custom_style(document_id: str, style_name: str, style_type: str, font_name: str = "宋体", font_size: int = 12, bold: bool = False, italic: bool = False, color: tuple = (0, 0, 0)) -> str:
    """
    创建自定义文本样式。
    
    Args:
        document_id: 文档的唯一标识符（必填）。
        style_name: 自定义样式的名称（必填）。
        style_type: 样式类型（'paragraph' 或 'character'，必填）。
        font_name: 字体名称（可选，默认：宋体）。
        font_size: 字号大小（可选，默认：12磅）。
        bold: 是否加粗（可选，默认：False）。
        italic: 是否斜体（可选，默认：False）。
        color: 字体颜色，RGB三元组（可选，默认：黑色(0,0,0)）。
    
    Returns:
        一个字符串，表示样式创建成功或失败的消息。
    
    示例:
        create_custom_style(document_id="doc1", style_name="强调", style_type="character", font_name="黑体", font_size=14, bold=True, color=(255,0,0))
    """
    # 参数验证
    if style_type.lower() not in ['paragraph', 'character']:
        raise ValueError(f"样式类型必须是'paragraph'或'character'，当前值: {style_type}")
    
    # 加载文档
    doc = Document(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    # 创建自定义样式
    if style_name in doc.styles:
        raise ValueError(f"样式 '{style_name}' 已存在。")
    
    if style_type.lower() == 'paragraph':
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor(*color)
    else:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor(*color)
    
    # 保存文档
    doc.save(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    return f"样式 '{style_name}' 已成功创建到文档 {document_id}。"

@mcp.tool()
def format_text(document_id: str, paragraph_index: int, start_pos: int, end_pos: int, style_name: str = "") -> str:
    """
    格式化指定文本区域。
    
    Args:
        document_id: 文档的唯一标识符（必填）。
        paragraph_index: 要格式化的段落索引（从0开始计数，必填）。
        start_pos: 要格式化的起始位置（从0开始计数，必填）。
        end_pos: 要格式化的结束位置（从0开始计数，含该位置，必填）。
        style_name: 要应用的样式名称（可选）。
    
    Returns:
        一个字符串，表示操作成功或失败的消息。
    
    示例:
        format_text(document_id="doc1", paragraph_index=2, start_pos=5, end_pos=10, style_name="强调")
    """
    # 加载文档
    doc = Document(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    # 验证段落索引
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise ValueError(f"段落索引超出范围: {paragraph_index}")
    
    paragraph = doc.paragraphs[paragraph_index]
    
    # 获取段落文本
    text = paragraph.text
    
    # 验证位置参数
    if start_pos < 0 or end_pos < 0 or start_pos > end_pos or end_pos >= len(text):
        raise ValueError(f"位置参数无效: start_pos={start_pos}, end_pos={end_pos}")
    
    # 分割文本
    before_text = text[:start_pos]
    formatted_text = text[start_pos:end_pos+1]
    after_text = text[end_pos+1:]
    
    # 清除原始段落
    paragraph.clear()
    
    # 添加未格式化部分
    run_before = paragraph.add_run(before_text)
    
    # 添加格式化部分
    run_formatted = paragraph.add_run(formatted_text)
    if style_name:
        try:
            run_formatted.style = doc.styles[style_name]
        except KeyError:
            raise ValueError(f"样式 '{style_name}' 不存在。")
    
    # 添加后续部分
    run_after = paragraph.add_run(after_text)
    
    # 保存文档
    doc.save(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    return f"文本已成功格式化到文档 {document_id}。"

@mcp.tool()
def protect_document(document_id: str, password: str) -> str:
    """
    设置文档密码保护。
    
    Args:
        document_id: 文档的唯一标识符（必填）。
        password: 用于保护文档的密码（必填）。
    
    Returns:
        一个字符串，表示操作成功或失败的消息。
    
    示例:
        protect_document(document_id="doc1", password="secret123")
    """
    # 加载文档
    doc = Document(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    # 设置文档保护
    try:
        # 尝试访问保护方法
        doc.protect(password=password)
    except Exception as e:
        # 如果文档保护功能不可用（例如，python-docx不支持直接设置文档保护）
        raise NotImplementedError("文档保护功能目前不可用，需要使用支持此功能的库或手动设置。") from e
    
    # 保存文档
    doc.save(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    return f"文档 {document_id} 已成功设置密码保护。"

@mcp.tool()
def add_footnote_to_document(document_id: str, paragraph_index: int, start_pos: int, end_pos: int, footnote_text: str, footnote_ref: str = "") -> str:
    """
    添加文档脚注。
    
    Args:
        document_id: 文档的唯一标识符（必填）。
        paragraph_index: 要添加脚注的段落索引（从0开始计数，必填）。
        start_pos: 要添加脚注的起始位置（从0开始计数，必填）。
        end_pos: 要添加脚注的结束位置（从0开始计数，含该位置，必填）。
        footnote_text: 脚注文本内容（必填）。
        footnote_ref: 脚注引用标记（可选，默认由系统自动分配）。
    
    Returns:
        一个字符串，表示操作成功或失败的消息。
    
    示例:
        add_footnote_to_document(document_id="doc1", paragraph_index=2, start_pos=5, end_pos=10, footnote_text="这是脚注内容")
    """
    # 加载文档
    doc = Document(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    # 验证段落索引
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise ValueError(f"段落索引超出范围: {paragraph_index}")
    
    paragraph = doc.paragraphs[paragraph_index]
    
    # 获取段落文本
    text = paragraph.text
    
    # 验证位置参数
    if start_pos < 0 or end_pos < 0 or start_pos > end_pos or end_pos >= len(text):
        raise ValueError(f"位置参数无效: start_pos={start_pos}, end_pos={end_pos}")
    
    # 分割文本
    before_text = text[:start_pos]
    annotated_text = text[start_pos:end_pos+1]
    after_text = text[end_pos+1:]
    
    # 清除原始段落
    paragraph.clear()
    
    # 添加未标注部分
    run_before = paragraph.add_run(before_text)
    
    # 添加标注部分
    run_annotated = paragraph.add_run(annotated_text)
    
    # 添加脚注
    try:
        # 尝试获取脚注部分
        footnotes_part = doc._part.footnotes_part
        if not footnotes_part:
            # 如果没有脚注部分，创建一个新的
            footnotes_part = doc._part._add_footnotes_part()
        
        # 添加脚注
        footnote = footnotes_part.add_footnote(footnote_text)
        
        # 添加脚注引用
        if footnote_ref:
            footnote.ref = footnote_ref
        
        # 在文档中添加脚注引用标记
        run_annotated.add_footnote_reference(footnote)
    except AttributeError as e:
        # 如果python-docx不支持直接添加脚注
        raise NotImplementedError("脚注功能目前不可用，需要使用支持此功能的扩展库或手动设置。") from e
    
    # 添加后续部分
    run_after = paragraph.add_run(after_text)
    
    # 保存文档
    doc.save(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    return f"脚注已成功添加到文档 {document_id}。"

@mcp.tool()
def get_paragraph_text_from_document(document_id: str, paragraph_index: int) -> str:
    """
    获取特定段落文本。
    
    Args:
        document_id: 文档的唯一标识符（必填）。
        paragraph_index: 要获取的段落索引（从0开始计数，必填）。
    
    Returns:
        一个字符串，包含请求的段落文本。
    
    示例:
        get_paragraph_text_from_document(document_id="doc1", paragraph_index=2)
    """
    # 加载文档
    doc = Document(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    # 验证段落索引
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise ValueError(f"段落索引超出范围: {paragraph_index}")
    
    # 获取段落文本
    text = doc.paragraphs[paragraph_index].text
    
    return text

@mcp.tool()
def find_text_in_document(document_id: str, search_text: str) -> list:
    """
    在文档中搜索指定文本。
    
    Args:
        document_id: 文档的唯一标识符（必填）。
        search_text: 要搜索的文本（必填）。
    
    Returns:
        一个列表，包含所有匹配的段落索引和文本内容。
    
    示例:
        find_text_in_document(document_id="doc1", search_text="关键词")
    """
    # 加载文档
    doc = Document(os.path.join(DOC_STORAGE_PATH, f"{document_id}.docx"))
    
    # 搜索文本
    results = []
    for i, paragraph in enumerate(doc.paragraphs):
        if search_text in paragraph.text:
            results.append({"paragraph_index": i, "text": paragraph.text})
    
    return results